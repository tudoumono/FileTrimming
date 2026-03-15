"""Step2: Word 構造抽出器

.docx ファイルから段落・表・図形を抽出し、中間表現 (IntermediateDocument) を構築する。

設計方針:
  - python-docx で確定的に抽出（LLM は使わない）
  - Oasys/Win スタイル対応: 太字が取れないため、フォントサイズ差と
    文字数ヒューリスティクスで疑似見出しを検出
  - 要素の出現順序を文書内の XML 順で保持
  - 変更履歴テーブル検出: 1行目に「ページ」「種別」「年月」「記事」のうち3個以上
"""

from __future__ import annotations

import re
import time
from logging import getLogger
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.config import PipelineConfig
from typing import Any

from src.models.intermediate import (
    CellData,
    Confidence,
    DocumentElement,
    ElementType,
    ImageElement,
    IntermediateDocument,
    ShapeElement,
)
from src.models.metadata import (
    ExtractedFileRecord,
    FileMetadata,
    ProcessStatus,
    StepResult,
)

logger = getLogger(__name__)

_FIGURE_CAPTION_RE = re.compile(
    r"^(?:図|Fig\.?|Figure)(?=\s*\d|\s*[:：]|\s+)(?:\s*[:：]?\s*)(.+)$",
    re.IGNORECASE,
)
_TABLE_CAPTION_RE = re.compile(
    r"^(?:表|Table)(?=\s*\d|\s*[:：]|\s+)(?:\s*[:：]?\s*)(.+)$",
    re.IGNORECASE,
)
_SECTION_NUMBER_RE = re.compile(
    r"^(?:第(\d+)章|(\d+(?:\.\d+)*)\.?\s)"
)
_ARROW_CHARS = set("→←↑↓⇒⇐⇑⇓▶▷►▸◀◁◄◂➔➡➜➞")


# ---------------------------------------------------------------------------
# 疑似見出し検出
# ---------------------------------------------------------------------------

def _get_font_size_pt(para: Paragraph) -> float | None:
    """段落の実効フォントサイズ (pt) を取得する。

    Run レベル → 段落スタイル → デフォルトの順で探索。
    """
    # Run レベルで最頻値を取る
    sizes: list[float] = []
    for run in para.runs:
        if run.font.size is not None:
            sizes.append(run.font.size.pt)
    if sizes:
        return max(set(sizes), key=sizes.count)

    # スタイルチェーン
    style = para.style
    while style is not None:
        if style.font and style.font.size is not None:
            return style.font.size.pt
        style = style.base_style
    return None


def _is_figure_caption(text: str) -> str | None:
    """段落テキストが図キャプションかどうか判定する。"""
    stripped = text.strip()
    if not stripped or len(stripped) > 60 or stripped.endswith("。"):
        return None

    match = _FIGURE_CAPTION_RE.match(stripped)
    if not match:
        return None

    caption = match.group(1).strip()
    return caption or None


def _is_table_caption(text: str) -> str | None:
    """段落テキストが表キャプションかどうか判定する。"""
    stripped = text.strip()
    if not stripped or len(stripped) > 60 or stripped.endswith("。"):
        return None

    match = _TABLE_CAPTION_RE.match(stripped)
    if not match:
        return None

    caption = match.group(1).strip()
    return caption or None


def _detect_section_number_depth(text: str) -> int | None:
    """テキスト先頭の section 番号から階層深度を検出する。"""
    match = _SECTION_NUMBER_RE.match(text.strip())
    if not match:
        return None

    if match.group(1):
        return 1
    if match.group(2):
        return match.group(2).count(".") + 1
    return None


def _is_arrow_annotation(text: str) -> bool:
    """矢印記号を含む注釈テキスト（フロー図の接続表現）か判定する。"""
    return any(char in _ARROW_CHARS for char in text)


def _detect_heading(
    para: Paragraph,
    config: PipelineConfig,
) -> tuple[int, str] | None:
    """疑似見出しを検出する。

    Returns:
        (level, detection_method) or None
    """
    text = para.text.strip()
    if not text or len(text) > config.heading_max_chars:
        return None

    # 1. Word 見出しスタイル（Heading 1, Heading 2 等）
    style_name = para.style.name or ""
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.split()[-1])
            return (level, "style")
        except (ValueError, IndexError):
            pass

    # 2. outlineLevel
    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        outline = pPr.find(qn("w:outlineLvl"))
        if outline is not None:
            val = outline.get(qn("w:val"))
            if val is not None:
                level = int(val) + 1  # outlineLevel は 0-indexed
                return (min(level, 6), "outline_level")

    # 3. フォントサイズ差（Oasys/Win 対応）
    size = _get_font_size_pt(para)
    if size is not None and size >= config.heading_font_size_min_pt:
        depth = _detect_section_number_depth(text)
        if depth is not None:
            level = min(depth + 1, 6)
        else:
            # サイズに応じてレベルを推定
            if size >= 16.0:
                level = 1
            elif size >= 14.0:
                level = 2
            elif size >= 12.0:
                level = 3
            else:
                level = 4
        return (level, f"font_size:{size}pt")

    # 4. 短文 + 行末句点なし（見出しらしいパターン）
    if len(text) <= 30 and not text.endswith(("。", ".", "、", ",")):
        # ただし数字のみ、空白のみは除外
        if re.search(r"[\u3040-\u9fff\uff01-\uff5ea-zA-Z]", text):
            if not _is_arrow_annotation(text):
                if _is_figure_caption(text) is not None:
                    return None
                if _is_table_caption(text) is not None:
                    return None

                depth = _detect_section_number_depth(text)
                if depth is not None:
                    level = min(depth + 1, 6)
                    return (level, "heuristic:section_number")

                return (3, "heuristic:short_no_period")

    return None


# ---------------------------------------------------------------------------
# 表の抽出
# ---------------------------------------------------------------------------

def _extract_table(table: Table) -> tuple[list[list[CellData]], bool]:
    """表からセルデータを抽出する。

    python-docx の row.cells は横結合セルを各列位置で同一 _tc として返す。
    id(tc) で重複を排除し、各セルを1回だけ追加する。

    Returns:
        (rows, has_merged_cells)
    """
    has_merged = False
    rows: list[list[CellData]] = []

    for r_idx, row in enumerate(table.rows):
        row_data: list[CellData] = []
        seen_tcs: set[int] = set()

        for c_idx, cell in enumerate(row.cells):
            tc = cell._tc
            tc_id = id(tc)

            # 横結合で同一 _tc が複数回返される → 重複スキップ
            if tc_id in seen_tcs:
                continue
            seen_tcs.add(tc_id)

            # 結合セルの検出（gridSpan / vMerge）
            rowspan = 1
            colspan = 1

            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is not None:
                gs = tc_pr.find(qn("w:gridSpan"))
                if gs is not None:
                    colspan = int(gs.get(qn("w:val"), "1"))
                vm = tc_pr.find(qn("w:vMerge"))
                if vm is not None:
                    val = vm.get(qn("w:val"), "")
                    if val != "restart":
                        # 継続セル（上のセルに結合されている）→ スキップ
                        continue

            if colspan > 1:
                has_merged = True

            # セルテキスト: 縦結合セルで複数段落が \n 連結され
            # 同一テキストが重複するケースの対策
            raw_parts = [p.strip() for p in cell.text.split("\n") if p.strip()]
            # 全パートが同一なら重複除去、それ以外はスペース結合
            if raw_parts and all(p == raw_parts[0] for p in raw_parts):
                cell_text = raw_parts[0]
            else:
                cell_text = " ".join(raw_parts)
            row_data.append(CellData(
                text=cell_text,
                row=r_idx,
                col=c_idx,
                rowspan=rowspan,
                colspan=colspan,
                is_header=(r_idx == 0),  # 1行目をヘッダーとみなす
            ))

        if row_data:
            rows.append(row_data)

    return rows, has_merged


def _is_change_history_table(
    rows: list[list[CellData]],
    config: PipelineConfig,
) -> bool:
    """変更履歴テーブルかどうかを判定する。

    1行目のセルテキストから全角スペースを除去した上で
    キーワードマッチングを行う。
    """
    if not rows or not rows[0]:
        return False

    matched = 0
    for cell in rows[0]:
        normalized = re.sub(r"[\s\u3000]+", "", cell.text)
        for kw in config.change_history_keywords:
            if kw in normalized:
                matched += 1
                break
    return matched >= config.change_history_min_match


# ---------------------------------------------------------------------------
# 図形の抽出
# ---------------------------------------------------------------------------

    # NOTE: _extract_shapes() は廃止。
    # 図形・画像は _build_element_order() → _has_floating_shape() / _has_inline_image()
    # で段落単位でインライン検出し、正しい出現位置に挿入するようにした。


# ---------------------------------------------------------------------------
# 文書要素の順序付き抽出
# ---------------------------------------------------------------------------

def _has_inline_image(para_elem) -> str | None:
    """段落要素内にインライン画像 (w:drawing 内の a:blip) があるか判定する。

    Returns:
        None: 画像なし
        str: 画像あり（alt text を返す。なければ空文字列）
    """
    ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    for drawing in para_elem.iter(qn("w:drawing")):
        # a:blip があれば画像
        has_blip = any(True for _ in drawing.iter(qn("a:blip")))
        if not has_blip:
            # pic:pic があれば画像
            ns_pic = "http://schemas.openxmlformats.org/drawingml/2006/picture"
            has_pic = any(True for _ in drawing.iter(f"{{{ns_pic}}}pic"))
            if not has_pic:
                continue
        # alt text を wp:docPr の descr 属性から取得
        alt_text = ""
        for doc_pr in drawing.iter(f"{{{ns_wp}}}docPr"):
            alt_text = doc_pr.get("descr", "")
            if not alt_text:
                alt_text = doc_pr.get("title", "")
            break
        return alt_text
    return None


def _has_vml_image(para_elem) -> str | None:
    """段落要素内に VML 画像 (w:pict 内の v:imagedata) があるか判定する。

    Returns:
        None: VML 画像なし
        str: VML 画像あり（alt text を返す。なければ空文字列）
    """
    for pict in para_elem.iter(qn("w:pict")):
        for child in pict.iter():
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "imagedata":
                # o:title 属性から alt text を取得
                alt_text = child.get("title", "") or child.get(
                    "{urn:schemas-microsoft-com:office:office}title", ""
                )
                return alt_text
    return None


def _parse_vml_style(style_str: str) -> dict[str, float]:
    """VML style 属性から位置・サイズ (pt) を抽出する。"""
    result: dict[str, float] = {}
    for part in style_str.split(";"):
        part = part.strip()
        if ":" not in part:
            continue
        key, val = part.split(":", 1)
        key = key.strip()
        val = val.strip()
        if key in ("left", "top", "width", "height"):
            m = re.match(r"([\d.]+)", val)
            if m:
                result[key] = float(m.group(1))
    return result


def _has_floating_shape(para_elem) -> list[ShapeElement]:
    """段落要素内の浮動図形・テキストボックスを抽出する。"""
    shapes: list[ShapeElement] = []

    # w:drawing 内の図形（画像以外）
    for drawing in para_elem.iter(qn("w:drawing")):
        # a:blip があれば画像なのでスキップ（_has_inline_image で処理）
        has_blip = any(True for _ in drawing.iter(qn("a:blip")))
        if has_blip:
            continue
        texts = [t.text for t in drawing.iter(qn("a:t")) if t.text]
        shapes.append(ShapeElement(
            shape_type="floating",
            texts=texts,
            confidence=Confidence.MEDIUM if texts else Confidence.LOW,
            fallback_reason="" if texts else "no_text_content",
        ))

    # VML 図形 (w:pict) — 画像以外
    for pict in para_elem.iter(qn("w:pict")):
        has_imagedata = False
        for child in pict.iter():
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "imagedata":
                has_imagedata = True
                break
        if has_imagedata:
            continue  # VML 画像はスキップ

        # VML 図形タイプと位置情報を取得
        texts = []
        shape_type = "vml"
        pos: dict[str, float] = {}
        for child in pict.iter():
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("shape", "rect", "oval", "line"):
                style_str = child.get("style", "")
                pos = _parse_vml_style(style_str)
                if tag == "rect":
                    shape_type = "vml_rect"
                elif tag == "shape":
                    # _x0000_t202 = テキストボックス
                    vtype = child.get("type", "")
                    if "_x0000_t202" in vtype:
                        shape_type = "vml_textbox"

        for t_elem in pict.iter():
            if t_elem.tag.endswith("}t") or t_elem.tag == "t":
                if t_elem.text:
                    texts.append(t_elem.text)

        shapes.append(ShapeElement(
            shape_type=shape_type,
            texts=texts,
            confidence=Confidence.MEDIUM if texts else Confidence.LOW,
            fallback_reason="" if texts else "no_text_content",
            left_pt=pos.get("left"),
            top_pt=pos.get("top"),
            width_pt=pos.get("width"),
            height_pt=pos.get("height"),
        ))

    return shapes


def _merge_overlapping_shapes(shapes: list[ShapeElement]) -> list[ShapeElement]:
    """重なる図形をマージする。

    テキストなし矩形 (vml_rect) とテキストあり図形 (vml_textbox 等) が
    同じ位置に重なっている場合、テキストなし矩形を除去する。
    """
    if len(shapes) < 2:
        return shapes

    # 位置情報を持つ図形を分類
    text_shapes: list[ShapeElement] = []    # テキストあり
    empty_rects: list[ShapeElement] = []    # テキストなし矩形
    no_pos: list[ShapeElement] = []         # 位置情報なし

    for s in shapes:
        if s.left_pt is None or s.top_pt is None:
            no_pos.append(s)
        elif not s.texts and s.shape_type == "vml_rect":
            empty_rects.append(s)
        else:
            text_shapes.append(s)

    # テキストなし矩形について、近くにテキストあり図形があれば除去
    OVERLAP_THRESHOLD_PT = 30.0  # 30pt 以内を「重なり」とみなす
    surviving_rects: list[ShapeElement] = []

    for rect in empty_rects:
        overlaps = False
        for ts in text_shapes:
            if ts.left_pt is None or ts.top_pt is None:
                continue
            dx = abs((rect.left_pt or 0) - (ts.left_pt or 0))
            dy = abs((rect.top_pt or 0) - (ts.top_pt or 0))
            if dx <= OVERLAP_THRESHOLD_PT and dy <= OVERLAP_THRESHOLD_PT:
                overlaps = True
                break
        if not overlaps:
            surviving_rects.append(rect)

    return no_pos + text_shapes + surviving_rects


def _group_shapes_as_flow(shapes: list[ShapeElement]) -> list[ShapeElement]:
    """連続する図形を位置情報でソートし、フロー図としてグルーピングする。"""
    if len(shapes) < 3:
        return shapes

    text_shapes = [shape for shape in shapes if shape.texts]
    if len(text_shapes) < 3:
        return shapes

    has_pos = [
        shape for shape in text_shapes
        if shape.top_pt is not None and shape.left_pt is not None
    ]
    no_pos = [
        shape for shape in text_shapes
        if shape.top_pt is None or shape.left_pt is None
    ]

    if len(has_pos) >= 2:
        has_pos.sort(key=lambda shape: (shape.top_pt or 0, shape.left_pt or 0))
        sorted_shapes = has_pos + no_pos
    else:
        sorted_shapes = text_shapes

    flow_texts: list[str] = []
    for shape in sorted_shapes:
        parts = [
            part.strip()
            for text in shape.texts
            for part in text.splitlines()
            if part.strip()
        ]
        combined = " / ".join(parts)
        if combined:
            flow_texts.append(combined)

    if not flow_texts:
        return shapes

    workflow = ShapeElement(
        shape_type="workflow",
        texts=flow_texts,
        confidence=Confidence.MEDIUM,
        fallback_reason="",
    )
    return [workflow]


def _para_has_own_text(para_elem) -> bool:
    """段落自体（図形・テキストボックス外）にテキストがあるか判定する。"""
    # pict / drawing 要素のセット（祖先チェック用）
    shape_roots: set = set(
        list(para_elem.iter(qn("w:pict"))) +
        list(para_elem.iter(qn("w:drawing")))
    )
    for t_elem in para_elem.iter(qn("w:t")):
        if not (t_elem.text and t_elem.text.strip()):
            continue
        # 祖先を辿って pict/drawing の中にいないか確認
        parent = t_elem.getparent()
        in_shape = False
        while parent is not None and parent is not para_elem:
            if parent in shape_roots:
                in_shape = True
                break
            parent = parent.getparent()
        if not in_shape:
            return True
    return False


def _build_element_order(doc: Document) -> list[tuple[str, Any]]:
    """文書本文の XML から要素の出現順序を構築する。

    Returns:
        [(element_type, data), ...]:
          - ("paragraph", para_idx)
          - ("table", table_idx)
          - ("image", alt_text_str)
          - ("shape_inline", ShapeElement)
    """
    order: list[tuple[str, Any]] = []
    para_idx = 0
    table_idx = 0

    # 図形のみの段落が続く場合、テキスト段落到達時にまとめて重なり判定するために蓄積
    pending_shapes: list[ShapeElement] = []

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            # 段落内の画像チェック（VML 画像より先に DrawingML を確認）
            inline_alt = _has_inline_image(child)
            if inline_alt is not None:
                order.append(("image", inline_alt))
            else:
                vml_alt = _has_vml_image(child)
                if vml_alt is not None:
                    order.append(("image", vml_alt))

            # 段落内の浮動図形チェック
            inline_shapes = _has_floating_shape(child)
            pending_shapes.extend(inline_shapes)

            # 図形外に自前のテキストがある段落に到達したら図形をフラッシュ
            if _para_has_own_text(child) and pending_shapes:
                merged = _merge_overlapping_shapes(pending_shapes)
                grouped = _group_shapes_as_flow(merged)
                for shape in grouped:
                    order.append(("shape_inline", shape))
                pending_shapes = []

            order.append(("paragraph", para_idx))
            para_idx += 1

        elif tag == "tbl":
            # テーブル前に蓄積された図形をフラッシュ
            if pending_shapes:
                merged = _merge_overlapping_shapes(pending_shapes)
                grouped = _group_shapes_as_flow(merged)
                for shape in grouped:
                    order.append(("shape_inline", shape))
                pending_shapes = []
            order.append(("table", table_idx))
            table_idx += 1

    # 末尾に残った図形をフラッシュ
    if pending_shapes:
        merged = _merge_overlapping_shapes(pending_shapes)
        grouped = _group_shapes_as_flow(merged)
        for shape in grouped:
            order.append(("shape_inline", shape))

    return order


# ---------------------------------------------------------------------------
# メインの抽出関数
# ---------------------------------------------------------------------------

def extract_docx(
    docx_path: Path,
    source_path: str,
    source_ext: str,
    config: PipelineConfig,
) -> tuple[ExtractedFileRecord, StepResult]:
    """1つの .docx ファイルから中間表現を抽出する。

    Args:
        docx_path: 正規化済み .docx のパス
        source_path: 元ファイルの相対パス (追跡用)
        source_ext: 元の拡張子
        config: パイプライン設定

    Returns:
        (ExtractedFileRecord, StepResult)
    """
    t0 = time.perf_counter()

    try:
        doc = Document(str(docx_path))
    except Exception as e:
        elapsed = time.perf_counter() - t0
        result = StepResult(
            file_path=source_path, step="extract",
            status=ProcessStatus.ERROR, message=str(e),
            duration_sec=round(elapsed, 2),
        )
        meta = FileMetadata(source_path=source_path, source_ext=source_ext)
        record = ExtractedFileRecord(metadata=meta, document={})
        return record, result

    intermediate = IntermediateDocument()

    # --- 表を先に抽出（テーブルオブジェクト → データ変換）---
    table_data_list: list[tuple[list[list[CellData]], bool, bool]] = []
    for table in doc.tables:
        rows, has_merged = _extract_table(table)
        is_ch = _is_change_history_table(rows, config)
        table_data_list.append((rows, has_merged, is_ch))

    # --- 要素を出現順に中間表現へ追加 ---
    element_order = _build_element_order(doc)
    paragraphs = doc.paragraphs
    source_idx = 0
    table_counter = 0
    shape_count = 0
    image_count = 0
    last_image_element: ImageElement | None = None

    for elem_type, idx_or_data in element_order:
        if elem_type == "image":
            # 段落内のインライン画像 → IMAGE 要素として追加
            alt_text = idx_or_data if isinstance(idx_or_data, str) else ""
            image = ImageElement(alt_text=alt_text, description="")
            intermediate.elements.append(DocumentElement(
                type=ElementType.IMAGE,
                content=image,
                source_index=source_idx,
            ))
            last_image_element = image
            source_idx += 1
            image_count += 1
            continue

        if elem_type == "shape_inline":
            # 段落内の浮動図形 → SHAPE 要素として追加（正しい出現位置）
            shape: ShapeElement = idx_or_data
            last_image_element = None
            shape_count += 1
            intermediate.elements.append(DocumentElement(
                type=ElementType.SHAPE,
                content=shape,
                source_index=source_idx,
            ))
            source_idx += 1
            continue

        if elem_type == "paragraph" and idx_or_data < len(paragraphs):
            para = paragraphs[idx_or_data]
            text = para.text.strip()

            if last_image_element is not None and text:
                caption_text = _is_figure_caption(text)
                if caption_text:
                    last_image_element.description = caption_text
                    last_image_element = None
                    source_idx += 1
                    continue

            if text:
                last_image_element = None

            # 疑似見出し検出
            heading_info = _detect_heading(para, config)
            if heading_info:
                level, method = heading_info
                intermediate.add_heading(level, text, method, source_index=source_idx)
            elif text:
                intermediate.add_paragraph(text, source_index=source_idx)

        elif elem_type == "table" and table_counter < len(table_data_list):
            last_image_element = None
            rows, has_merged, is_ch = table_data_list[table_counter]
            table_counter += 1

            # 表の直前段落をキャプション候補として取得
            caption = ""
            if intermediate.elements:
                last = intermediate.elements[-1]
                if last.type.value == "paragraph" and last.content is not None:
                    candidate = last.content.text  # type: ignore[union-attr]
                    if len(candidate) <= 60:
                        caption = candidate
                        intermediate.elements.pop()

            confidence = Confidence.HIGH
            fallback_reason = ""
            if has_merged:
                confidence = Confidence.MEDIUM
            if is_ch:
                fallback_reason = "change_history_table"

            intermediate.add_table(
                rows=rows,
                caption=caption,
                has_merged_cells=has_merged,
                confidence=confidence,
                fallback_reason=fallback_reason,
                source_index=source_idx,
            )
        else:
            last_image_element = None

        source_idx += 1

    # --- 変更履歴テーブル数から doc_role 推定 ---
    total_tables = len(table_data_list)
    ch_count = sum(1 for _, _, is_ch in table_data_list if is_ch)
    if total_tables == 0:
        doc_role = "unknown"
    elif ch_count > 0 and ch_count == total_tables:
        doc_role = "change_history"
    elif ch_count > 0:
        doc_role = "mixed"
    else:
        doc_role = "spec_body"

    # --- メタデータ構築 ---
    meta = FileMetadata(
        source_path=source_path,
        source_ext=source_ext,
        source_size_bytes=docx_path.stat().st_size,
        normalized_from=source_ext if source_ext != ".docx" else "",
        doc_role_guess=doc_role,
    )

    record = ExtractedFileRecord(
        metadata=meta,
        document=intermediate.to_dict(),
    )

    elapsed = time.perf_counter() - t0
    warnings: list[str] = []
    if ch_count > 0:
        warnings.append(f"change_history_tables={ch_count}")
    if shape_count > 0:
        warnings.append(f"shapes={shape_count}")
    if image_count > 0:
        warnings.append(f"images={image_count}")

    status = ProcessStatus.SUCCESS
    msg = f"elements={len(intermediate.elements)}, tables={total_tables}, doc_role={doc_role}"
    if warnings:
        status = ProcessStatus.WARNING
        msg += ", " + ", ".join(warnings)

    result = StepResult(
        file_path=source_path, step="extract",
        status=status, message=msg,
        duration_sec=round(elapsed, 2),
    )
    logger.info("抽出完了: %s (%s)", source_path, msg)

    return record, result
