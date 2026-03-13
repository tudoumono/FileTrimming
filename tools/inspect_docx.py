"""docx 詳細調査スクリプト

.docx ファイルの内部構造を詳細に調査し、パイプライン設計に必要な
情報を抽出する。profile_documents.py より深い情報を取る。

調査項目:
  1. フォントパターン: 段落ごとのフォントサイズ・太字・スタイル名 → 疑似見出し検出の材料
  2. 表の詳細構造: 結合セル位置、列見出し/行見出し、入れ子表、親子行パターン
  3. 図形の詳細: 浮動図形のテキスト・座標・接続関係、テキストボックス
  4. 段落と表/図形の出現順序: 文書内の要素配置パターン

使い方:
  # 単一ファイルを調査
  python tools/inspect_docx.py <file.docx> -o inspect_report

  # フォルダ配下の全 .docx を調査
  python tools/inspect_docx.py <folder> -o inspect_report

出力:
  inspect_report.json  — 詳細データ
  inspect_report.txt   — 口頭説明用テキスト
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from dataclasses import asdict, dataclass, field
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("エラー: python-docx がインストールされていません (pip install python-docx)", file=sys.stderr)
    sys.exit(1)

try:
    from lxml import etree

    HAS_LXML = True
except ImportError:
    HAS_LXML = False


# ---------------------------------------------------------------------------
# データクラス
# ---------------------------------------------------------------------------
@dataclass
class ParagraphInfo:
    index: int
    text_preview: str  # 先頭50字
    char_count: int
    style_name: str
    font_size_pt: float | None
    is_bold: bool
    font_name: str | None
    is_all_caps: bool
    outline_level: int | None  # Word の outlineLevel（見出しレベルの別手段）
    pseudo_heading_reason: str = ""  # 疑似見出し判定理由
    bold_debug: str = ""  # 太字判定のデバッグ情報


@dataclass
class CellDetail:
    row: int
    col: int
    text_preview: str
    is_merged_horizontally: bool
    is_merged_vertically: bool
    grid_span: int  # 横方向の結合数


@dataclass
class TableDetail:
    table_index: int
    rows: int
    cols: int
    has_merged_cells: bool
    merged_cell_count: int
    first_row_texts: list[str]  # 1行目のテキスト（列見出しの推定）
    first_col_texts: list[str]  # 1列目のテキスト（行見出しの推定）
    has_nested_table: bool
    max_cell_text_length: int
    sample_cells: list[dict]  # 結合セルなど注目すべきセルのサンプル


@dataclass
class ShapeDetail:
    shape_type: str  # "floating" / "inline" / "textbox" / "group"
    text: str | None
    width: float | None  # EMU → cm
    height: float | None
    left: float | None  # アンカー位置
    top: float | None
    has_text_content: bool
    child_count: int  # グループ内の子要素数


@dataclass
class ElementOrder:
    """文書内の要素出現順序を記録する。"""

    index: int
    element_type: str  # "paragraph" / "table" / "shape"
    preview: str  # 内容のプレビュー
    is_pseudo_heading: bool  # フォントパターンから疑似見出しと判定されたか


@dataclass
class DocxInspection:
    path: str
    error: str | None = None

    # フォントパターン
    paragraphs: list[dict] = field(default_factory=list)
    font_size_distribution: dict[str, int] = field(default_factory=dict)
    bold_paragraph_count: int = 0
    pseudo_heading_candidates: list[dict] = field(default_factory=list)

    # 表の詳細
    tables: list[dict] = field(default_factory=list)
    total_merged_cells: int = 0
    nested_table_count: int = 0

    # 図形の詳細
    shapes: list[dict] = field(default_factory=list)
    shapes_with_text: int = 0

    # 要素の出現順序
    element_order: list[dict] = field(default_factory=list)

    # 変更履歴テーブル検出
    change_history_tables: list[dict] = field(default_factory=list)  # 検出された変更履歴テーブルの情報
    doc_role_guess: str = "unknown"  # "change_history" / "spec_body" / "mixed" / "unknown"


# ---------------------------------------------------------------------------
# フォントサイズ取得ヘルパー
# ---------------------------------------------------------------------------
def get_font_size_pt(para) -> float | None:
    """段落のフォントサイズを pt で取得する。"""
    # 段落スタイルのフォントサイズ
    if para.style and para.style.font and para.style.font.size:
        return para.style.font.size.pt

    # Run レベルのフォントサイズ（最初の Run を参照）
    for run in para.runs:
        if run.font.size:
            return run.font.size.pt

    # XML から直接取得
    rpr = para._element.find(qn("w:pPr"))
    if rpr is not None:
        rfonts = rpr.find(qn("w:rPr"))
        if rfonts is not None:
            sz = rfonts.find(qn("w:sz"))
            if sz is not None:
                val = sz.get(qn("w:val"))
                if val:
                    return int(val) / 2  # half-points → points

    return None


def _resolve_style_bold(style) -> bool | None:
    """スタイルの継承チェーンを辿って太字かどうかを解決する。"""
    visited = set()
    current = style
    while current and current.style_id not in visited:
        visited.add(current.style_id)
        # スタイル定義の XML から w:b を探す
        style_elem = current._element
        rpr = style_elem.find(qn("w:rPr"))
        if rpr is not None:
            b = rpr.find(qn("w:b"))
            if b is not None:
                val = b.get(qn("w:val"))
                if val is not None and val.lower() in ("false", "0"):
                    return False
                return True  # w:b が存在すれば太字（val なし or true/1）
        # python-docx の font.bold が明示的に設定されていれば使う
        if current.font and current.font.bold is not None:
            return current.font.bold
        # 親スタイルへ辿る
        current = current.base_style
    return None


def get_is_bold(para) -> tuple[bool, str]:
    """段落が太字かどうかを判定する。デバッグ情報も返す。"""
    # 1. 段落スタイルの継承チェーンから太字を解決
    if para.style:
        style_bold = _resolve_style_bold(para.style)
        if style_bold is True:
            return True, f"style_chain({para.style.name})"

    # 2. 段落レベルの太字設定（pPr/rPr/b）
    ppr = para._element.find(qn("w:pPr"))
    if ppr is not None:
        rpr = ppr.find(qn("w:rPr"))
        if rpr is not None:
            b = rpr.find(qn("w:b"))
            if b is not None:
                val = b.get(qn("w:val"))
                if val is None or val.lower() in ("true", "1", ""):
                    return True, "pPr/rPr/b"

    # 3. Run レベルで判定
    runs_with_text = [run for run in para.runs if run.text.strip()]
    if not runs_with_text:
        return False, f"runs=0/total_runs={len(para.runs)}"
    bold_count = 0
    xml_bold_count = 0
    none_count = 0
    false_count = 0
    for run in runs_with_text:
        if run.bold:
            bold_count += 1
        elif run.bold is None:
            none_count += 1
            # Run のスタイル（rStyle）経由の太字チェック
            run_rpr = run._element.find(qn("w:rPr"))
            if run_rpr is not None:
                b = run_rpr.find(qn("w:b"))
                if b is not None:
                    val = b.get(qn("w:val"))
                    if val is None or val.lower() in ("true", "1", ""):
                        xml_bold_count += 1
                        bold_count += 1
                        continue
                # rStyle 参照があればそのスタイルの太字もチェック
                rstyle = run_rpr.find(qn("w:rStyle"))
                if rstyle is not None:
                    rstyle_val = rstyle.get(qn("w:val"))
                    if rstyle_val and run.part and run.part.document:
                        try:
                            char_style = run.part.document.styles[rstyle_val]
                            if _resolve_style_bold(char_style):
                                xml_bold_count += 1
                                bold_count += 1
                                continue
                        except KeyError:
                            pass
        else:
            false_count += 1
    total = len(runs_with_text)
    is_bold = bold_count > total / 2
    debug = f"runs={total},bold={bold_count - xml_bold_count},xml_b={xml_bold_count},none={none_count},false={false_count}"
    return is_bold, debug


def get_font_name(para) -> str | None:
    """段落のフォント名を取得する。"""
    for run in para.runs:
        if run.font.name:
            return run.font.name
    return None


def get_is_all_caps(para) -> bool:
    """段落が全角大文字（all caps）かどうか。"""
    for run in para.runs:
        if run.font.all_caps:
            return True
    return False


def get_outline_level(para) -> int | None:
    """段落の outlineLevel を取得する（見出しスタイルとは別の見出し判定手段）。"""
    ppr = para._element.find(qn("w:pPr"))
    if ppr is not None:
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is not None:
            val = outline.get(qn("w:val"))
            if val is not None:
                return int(val)
    return None


# ---------------------------------------------------------------------------
# 表の詳細調査
# ---------------------------------------------------------------------------
def inspect_table(tbl, table_index: int) -> TableDetail:
    rows = len(tbl.rows)
    cols = len(tbl.columns)

    merged_cell_count = 0
    has_nested = False
    max_cell_text = 0
    sample_cells: list[dict] = []

    for r_idx, row in enumerate(tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            max_cell_text = max(max_cell_text, len(text))

            # 結合セル検出
            tc = cell._tc
            grid_span = 1
            gs = tc.find(qn("w:tcPr"))
            is_h_merged = False
            is_v_merged = False

            if gs is not None:
                span_elem = gs.find(qn("w:gridSpan"))
                if span_elem is not None:
                    val = span_elem.get(qn("w:val"))
                    if val and int(val) > 1:
                        grid_span = int(val)
                        is_h_merged = True

                vmerge = gs.find(qn("w:vMerge"))
                if vmerge is not None:
                    is_v_merged = True

            if is_h_merged or is_v_merged:
                merged_cell_count += 1
                if len(sample_cells) < 10:
                    sample_cells.append(
                        asdict(
                            CellDetail(
                                row=r_idx,
                                col=c_idx,
                                text_preview=text[:50],
                                is_merged_horizontally=is_h_merged,
                                is_merged_vertically=is_v_merged,
                                grid_span=grid_span,
                            )
                        )
                    )

            # 入れ子表検出
            nested = tc.findall(qn("w:tbl"))
            if nested:
                has_nested = True

    # 1行目・1列目のテキスト（見出し推定用）
    first_row_texts = []
    if rows > 0:
        for cell in tbl.rows[0].cells:
            first_row_texts.append(cell.text.strip()[:30])

    first_col_texts = []
    for row in tbl.rows[:10]:  # 先頭10行分
        if row.cells:
            first_col_texts.append(row.cells[0].text.strip()[:30])

    actual_cells = sum(1 for row in tbl.rows for _ in row.cells)
    grid_cells = rows * cols
    has_merged = actual_cells != grid_cells or merged_cell_count > 0

    return TableDetail(
        table_index=table_index,
        rows=rows,
        cols=cols,
        has_merged_cells=has_merged,
        merged_cell_count=merged_cell_count,
        first_row_texts=first_row_texts,
        first_col_texts=first_col_texts,
        has_nested_table=has_nested,
        max_cell_text_length=max_cell_text,
        sample_cells=sample_cells,
    )


# ---------------------------------------------------------------------------
# 図形の詳細調査
# ---------------------------------------------------------------------------
def inspect_shapes(doc) -> list[ShapeDetail]:
    shapes: list[ShapeDetail] = []

    # InlineShape
    for ishape in doc.inline_shapes:
        shapes.append(
            ShapeDetail(
                shape_type="inline",
                text=None,
                width=ishape.width / 914400 if ishape.width else None,  # EMU → cm (approx)
                height=ishape.height / 914400 if ishape.height else None,
                left=None,
                top=None,
                has_text_content=False,
                child_count=0,
            )
        )

    if not HAS_LXML:
        return shapes

    # 浮動図形を XML から取得
    body = doc.element.body
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "v": "urn:schemas-microsoft-com:vml",
        "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
    }

    # wps:wsp (WordprocessingShape)
    for wsp in body.findall(".//wps:wsp", ns):
        text_parts = []
        for txbx in wsp.findall(".//wps:txbx//w:p", ns):
            t = "".join(node.text or "" for node in txbx.findall(".//w:t", ns))
            if t.strip():
                text_parts.append(t.strip())

        text = "\n".join(text_parts) if text_parts else None
        shapes.append(
            ShapeDetail(
                shape_type="floating",
                text=text[:200] if text else None,
                width=None,
                height=None,
                left=None,
                top=None,
                has_text_content=bool(text),
                child_count=0,
            )
        )

    # VML shapes (v:shape)
    for vshape in body.findall(".//v:shape", ns):
        text_parts = []
        for textbox in vshape.findall(".//v:textbox//w:p", ns):
            t = "".join(node.text or "" for node in textbox.findall(".//w:t", ns))
            if t.strip():
                text_parts.append(t.strip())

        text = "\n".join(text_parts) if text_parts else None
        style = vshape.get("style", "")

        # style から幅・高さを抽出
        width = None
        height = None
        for part in style.split(";"):
            part = part.strip()
            if part.startswith("width:"):
                try:
                    width = float(part.replace("width:", "").replace("pt", "").strip()) / 72 * 2.54
                except ValueError:
                    pass
            elif part.startswith("height:"):
                try:
                    height = float(part.replace("height:", "").replace("pt", "").strip()) / 72 * 2.54
                except ValueError:
                    pass

        shapes.append(
            ShapeDetail(
                shape_type="vml",
                text=text[:200] if text else None,
                width=width,
                height=height,
                left=None,
                top=None,
                has_text_content=bool(text),
                child_count=0,
            )
        )

    # グループ図形 (wpg:wgp)
    for wgp in body.findall(".//wpg:wgp", ns):
        child_shapes = wgp.findall(".//wps:wsp", ns)
        text_parts = []
        for child in child_shapes:
            for txbx in child.findall(".//wps:txbx//w:p", ns):
                t = "".join(node.text or "" for node in txbx.findall(".//w:t", ns))
                if t.strip():
                    text_parts.append(t.strip())

        text = "\n".join(text_parts) if text_parts else None
        shapes.append(
            ShapeDetail(
                shape_type="group",
                text=text[:200] if text else None,
                width=None,
                height=None,
                left=None,
                top=None,
                has_text_content=bool(text),
                child_count=len(child_shapes),
            )
        )

    return shapes


# ---------------------------------------------------------------------------
# 疑似見出し候補の判定
# ---------------------------------------------------------------------------
def detect_pseudo_headings(paragraphs: list[ParagraphInfo]) -> list[ParagraphInfo]:
    """フォントパターンから疑似見出し候補を検出する。"""
    if not paragraphs:
        return []

    # 本文のフォントサイズを推定（最頻値）
    sizes = [p.font_size_pt for p in paragraphs if p.font_size_pt and p.char_count > 0]
    if not sizes:
        # サイズが取れない場合は太字のみで判定
        return [
            p
            for p in paragraphs
            if p.is_bold and 0 < p.char_count <= 100
        ]

    size_counter = Counter(sizes)
    body_size = size_counter.most_common(1)[0][0] if size_counter else 0

    candidates = []
    for p in paragraphs:
        if p.char_count == 0:
            continue

        is_candidate = False
        reason = ""

        # フォントサイズが本文より大きい
        if p.font_size_pt and p.font_size_pt > body_size:
            is_candidate = True
            reason = f"フォントサイズ大({p.font_size_pt}pt > 本文{body_size}pt)"
        # 本文と同じサイズで太字、かつ短い段落
        elif p.is_bold and p.char_count <= 100:
            is_candidate = True
            reason = f"太字+短文({p.char_count}字)"
        # 太字でフォントサイズが未取得（スタイル経由等）、かつ短い段落
        elif p.is_bold and p.font_size_pt is None and p.char_count <= 100:
            is_candidate = True
            reason = "太字+サイズ不明+短文"
        # outlineLevel が設定されている
        elif p.outline_level is not None:
            is_candidate = True
            reason = f"outlineLevel={p.outline_level}"

        if is_candidate:
            p.pseudo_heading_reason = reason
            candidates.append(p)

    return candidates


# ---------------------------------------------------------------------------
# 要素の出現順序を取得
# ---------------------------------------------------------------------------
def get_element_order(doc, pseudo_headings: list[ParagraphInfo]) -> list[ElementOrder]:
    """段落・表・図形の出現順序を記録する。"""
    pseudo_heading_indices = {p.index for p in pseudo_headings}
    elements: list[ElementOrder] = []
    idx = 0

    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            text = child.text or ""
            # w:t からテキストを集める
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            text_nodes = child.findall(".//w:t", ns)
            text = "".join(t.text or "" for t in text_nodes)

            elements.append(
                ElementOrder(
                    index=idx,
                    element_type="paragraph",
                    preview=text[:50] if text else "(空段落)",
                    is_pseudo_heading=idx in pseudo_heading_indices,
                )
            )
        elif tag == "tbl":
            # 表の最初のセルのテキストをプレビュー
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            first_texts = []
            for tc in child.findall(".//w:tc", ns)[:3]:
                t_nodes = tc.findall(".//w:t", ns)
                t = "".join(n.text or "" for n in t_nodes)
                if t.strip():
                    first_texts.append(t.strip()[:20])
            preview = " | ".join(first_texts) if first_texts else "(空表)"
            elements.append(
                ElementOrder(
                    index=idx,
                    element_type="table",
                    preview=preview,
                    is_pseudo_heading=False,
                )
            )
        else:
            continue

        idx += 1

    return elements


# ---------------------------------------------------------------------------
# 変更履歴テーブル検出
# ---------------------------------------------------------------------------
# 変更履歴テーブルの1行目に含まれるキーワード群
CHANGE_HISTORY_KEYWORDS = {"ページ", "種別", "年月", "記事"}
# 最低何個のキーワードが一致すれば変更履歴とみなすか
CHANGE_HISTORY_MIN_MATCH = 3


def detect_change_history_table(table_detail: dict) -> dict | None:
    """表の1行目テキストから変更履歴テーブルかどうかを判定する。

    Returns:
        検出情報の dict、または該当しなければ None。
    """
    first_row = table_detail.get("first_row_texts", [])
    if not first_row:
        return None

    # 1行目の各セルテキストとキーワードの一致をチェック
    # 全角スペース・半角スペース・タブなどを除去してから照合
    matched_keywords: list[str] = []
    for cell_text in first_row:
        normalized = re.sub(r"[\s\u3000]+", "", cell_text)
        for kw in CHANGE_HISTORY_KEYWORDS:
            if kw in normalized:
                matched_keywords.append(kw)
                break  # 1セルにつき1キーワードまで

    if len(matched_keywords) >= CHANGE_HISTORY_MIN_MATCH:
        return {
            "table_index": table_detail["table_index"],
            "rows": table_detail["rows"],
            "cols": table_detail["cols"],
            "first_row_texts": first_row,
            "matched_keywords": matched_keywords,
            "match_count": len(matched_keywords),
        }
    return None


# ---------------------------------------------------------------------------
# 1ファイルの調査
# ---------------------------------------------------------------------------
def inspect_file(filepath: Path) -> DocxInspection:
    rel_path = str(filepath)
    inspection = DocxInspection(path=rel_path)

    try:
        doc = Document(str(filepath))
    except Exception as e:
        inspection.error = f"読み込み失敗: {e}"
        return inspection

    # --- 段落のフォントパターン（一時コメントアウト: 太字検出の調査中） ---
    # para_infos: list[ParagraphInfo] = []
    # font_sizes: list[str] = []
    #
    # for i, para in enumerate(doc.paragraphs):
    #     text = para.text.strip()
    #     font_size = get_font_size_pt(para)
    #     is_bold, bold_debug = get_is_bold(para)
    #     font_name = get_font_name(para)
    #     is_all_caps = get_is_all_caps(para)
    #     outline_level = get_outline_level(para)
    #     style_name = para.style.name if para.style else ""
    #
    #     info = ParagraphInfo(
    #         index=i,
    #         text_preview=text[:50] if text else "",
    #         char_count=len(text),
    #         style_name=style_name,
    #         font_size_pt=font_size,
    #         is_bold=is_bold,
    #         font_name=font_name,
    #         is_all_caps=is_all_caps,
    #         outline_level=outline_level,
    #         bold_debug=bold_debug,
    #     )
    #     para_infos.append(info)
    #
    #     if font_size:
    #         font_sizes.append(f"{font_size}pt")
    #
    # inspection.paragraphs = [asdict(p) for p in para_infos]
    # inspection.font_size_distribution = dict(Counter(font_sizes).most_common())
    # inspection.bold_paragraph_count = sum(1 for p in para_infos if p.is_bold)
    #
    # # 疑似見出し検出
    # pseudo_headings = detect_pseudo_headings(para_infos)
    # inspection.pseudo_heading_candidates = [asdict(p) for p in pseudo_headings]

    # --- 表の詳細 ---
    for t_idx, tbl in enumerate(doc.tables):
        detail = inspect_table(tbl, t_idx)
        inspection.tables.append(asdict(detail))
        inspection.total_merged_cells += detail.merged_cell_count
        if detail.has_nested_table:
            inspection.nested_table_count += 1

    # --- 変更履歴テーブル検出 ---
    for t_dict in inspection.tables:
        ch_info = detect_change_history_table(t_dict)
        if ch_info:
            ch_info["_file"] = rel_path
            inspection.change_history_tables.append(ch_info)

    # 文書の役割推定
    total_tables = len(inspection.tables)
    ch_count = len(inspection.change_history_tables)
    if total_tables == 0:
        inspection.doc_role_guess = "unknown"
    elif ch_count > 0 and ch_count == total_tables:
        inspection.doc_role_guess = "change_history"
    elif ch_count > 0 and ch_count < total_tables:
        inspection.doc_role_guess = "mixed"
    elif ch_count == 0:
        inspection.doc_role_guess = "spec_body"

    # --- 図形の詳細（一時コメントアウト） ---
    # shape_details = inspect_shapes(doc)
    # inspection.shapes = [asdict(s) for s in shape_details]
    # inspection.shapes_with_text = sum(1 for s in shape_details if s.has_text_content)

    # --- 要素の出現順序（一時コメントアウト） ---
    # order = get_element_order(doc, pseudo_headings)
    # inspection.element_order = [asdict(e) for e in order]

    return inspection


# ---------------------------------------------------------------------------
# テキストレポート生成
# ---------------------------------------------------------------------------
def format_size(size_bytes: int) -> str:
    if size_bytes >= 1024 * 1024:
        return f"{size_bytes / 1024 / 1024:.1f}MB"
    elif size_bytes >= 1024:
        return f"{size_bytes / 1024:.0f}KB"
    return f"{size_bytes}バイト"


def build_text_report(inspections: list[DocxInspection]) -> str:
    lines: list[str] = []
    lines.append("=" * 60)
    lines.append("docx 詳細調査レポート")
    lines.append("=" * 60)
    lines.append("")
    lines.append(f"調査対象: {len(inspections)}ファイル")
    errors = [i for i in inspections if i.error]
    if errors:
        lines.append(f"エラー: {len(errors)}件")
    lines.append("")

    ok_files = [i for i in inspections if not i.error]
    if not ok_files:
        lines.append("調査可能なファイルがありません。")
        return "\n".join(lines)

    # --- フォントパターン集計（一時コメントアウト: 太字検出の調査中） ---
    # lines.append("■ フォントパターン分析")
    # lines.append("")
    #
    # all_sizes: Counter = Counter()
    # for insp in ok_files:
    #     for size_str, count in insp.font_size_distribution.items():
    #         all_sizes[size_str] += count
    #
    # if all_sizes:
    #     lines.append("  フォントサイズの分布（全ファイル合算）:")
    #     for size_str, count in all_sizes.most_common():
    #         lines.append(f"    {size_str}: {count}段落")
    #     body_size = all_sizes.most_common(1)[0][0]
    #     lines.append(f"  → 本文の推定フォントサイズ: {body_size}（最頻値）")
    # else:
    #     lines.append("  フォントサイズ: 取得できず（スタイル経由の可能性）")
    # lines.append("")
    #
    # total_bold = sum(i.bold_paragraph_count for i in ok_files)
    # total_paras = sum(len(i.paragraphs) for i in ok_files)
    # lines.append(f"  太字の段落: {total_bold}個 / {total_paras}段落中")
    #
    # bold_debug_counter: Counter = Counter()
    # for insp in ok_files:
    #     for p in insp.paragraphs:
    #         debug = p.get("bold_debug", "")
    #         if debug:
    #             bold_debug_counter[debug] += 1
    # if bold_debug_counter:
    #     lines.append("  [デバッグ] 太字判定の内訳（上位10パターン）:")
    #     for pattern, count in bold_debug_counter.most_common(10):
    #         lines.append(f"    {pattern}: {count}段落")
    #     none_samples = []
    #     for insp in ok_files:
    #         for p in insp.paragraphs:
    #             if p.get("char_count", 0) > 0 and "none=" in p.get("bold_debug", "") and len(none_samples) < 5:
    #                 none_samples.append(p)
    #     if none_samples:
    #         lines.append("  [デバッグ] run.bold=None のサンプル段落:")
    #         for s in none_samples:
    #             lines.append(f"    「{s['text_preview']}」 style={s['style_name']} {s['bold_debug']}")
    # lines.append("")
    #
    # total_pseudo = sum(len(i.pseudo_heading_candidates) for i in ok_files)
    # files_with_pseudo = sum(1 for i in ok_files if i.pseudo_heading_candidates)
    # lines.append(f"  疑似見出し候補: 合計{total_pseudo}個（{files_with_pseudo}ファイルで検出）")
    # if total_pseudo > 0:
    #     reason_counter: Counter = Counter()
    #     for insp in ok_files:
    #         for ph in insp.pseudo_heading_candidates:
    #             reason = ph.get("pseudo_heading_reason", "不明")
    #             reason_counter[reason] += 1
    #     lines.append("  判定理由の内訳:")
    #     for reason, count in reason_counter.most_common():
    #         lines.append(f"    {reason}: {count}個")
    #     lines.append("")
    #
    #     pseudo_texts: Counter = Counter()
    #     for insp in ok_files:
    #         for ph in insp.pseudo_heading_candidates:
    #             text = ph.get("text_preview", "")
    #             if text:
    #                 pseudo_texts[text] += 1
    #     lines.append("  疑似見出しとして検出されたテキスト（出現頻度順、上位20件）:")
    #     for text, count in pseudo_texts.most_common(20):
    #         lines.append(f"    「{text}」: {count}件")
    # lines.append("")
    #
    # has_outline = 0
    # for insp in ok_files:
    #     for p in insp.paragraphs:
    #         if p.get("outline_level") is not None:
    #             has_outline += 1
    # if has_outline > 0:
    #     lines.append(f"  outlineLevel が設定された段落: {has_outline}個")
    #     lines.append("  → Word の見出しスタイルは未使用だが outlineLevel で構造化されている可能性あり")
    # else:
    #     lines.append("  outlineLevel: 設定なし（見出し構造は完全にフォント依存）")
    # lines.append("")

    # --- 表の詳細集計 ---
    lines.append("■ 表の詳細分析")
    lines.append("")

    all_tables = []
    for insp in ok_files:
        for t in insp.tables:
            t["_file"] = insp.path
            all_tables.append(t)

    lines.append(f"  表の総数: {len(all_tables)}個")

    # 結合セル
    tables_with_merge = [t for t in all_tables if t["has_merged_cells"]]
    total_merged = sum(t["merged_cell_count"] for t in all_tables)
    lines.append(f"  結合セルを含む表: {len(tables_with_merge)}個（結合セル合計{total_merged}個）")

    if tables_with_merge:
        lines.append("  結合セルの詳細（サンプル）:")
        shown = 0
        for t in tables_with_merge:
            if shown >= 5:
                lines.append(f"    ... 他{len(tables_with_merge) - 5}個")
                break
            lines.append(f"    - {t['_file']} / 表{t['table_index']}: {t['rows']}行×{t['cols']}列、結合{t['merged_cell_count']}箇所")
            for sc in t.get("sample_cells", [])[:3]:
                merge_type = []
                if sc.get("is_merged_horizontally"):
                    merge_type.append(f"横{sc['grid_span']}セル結合")
                if sc.get("is_merged_vertically"):
                    merge_type.append("縦結合")
                lines.append(f"      ({sc['row']},{sc['col']}): {', '.join(merge_type)} 「{sc['text_preview']}」")
            shown += 1
    lines.append("")

    # 入れ子表
    nested = [t for t in all_tables if t["has_nested_table"]]
    lines.append(f"  入れ子表（表の中に表）: {len(nested)}個")
    if nested:
        for t in nested[:3]:
            lines.append(f"    - {t['_file']} / 表{t['table_index']}")
    lines.append("")

    # 列見出しパターン
    first_row_patterns: Counter = Counter()
    for t in all_tables:
        if t["first_row_texts"]:
            key = " | ".join(t["first_row_texts"][:5])
            first_row_patterns[key] += 1

    if first_row_patterns:
        lines.append("  表の1行目パターン（列見出しの推定、出現頻度順）:")
        for pattern, count in first_row_patterns.most_common(10):
            lines.append(f"    [{count}回] {pattern}")
    lines.append("")

    # 大きなセル
    large_cell_tables = [t for t in all_tables if t["max_cell_text_length"] >= 100]
    if large_cell_tables:
        lines.append(f"  100字以上のセルを含む表: {len(large_cell_tables)}個")
        lines.append("  → セル内に長文があり、チャンキング時の考慮が必要")
    lines.append("")

    # --- 変更履歴テーブル検出結果 ---
    lines.append("■ 変更履歴テーブル検出")
    lines.append("")

    # 文書の役割分類
    role_counter: Counter = Counter()
    for insp in ok_files:
        role_counter[insp.doc_role_guess] += 1

    role_labels = {
        "change_history": "変更履歴のみ",
        "spec_body": "仕様書本体（変更履歴なし）",
        "mixed": "仕様書+変更履歴の混在",
        "unknown": "不明（表なし等）",
    }
    lines.append("  文書の役割推定:")
    for role, count in role_counter.most_common():
        label = role_labels.get(role, role)
        lines.append(f"    {label}: {count}ファイル")
    lines.append("")

    # 検出された変更履歴テーブルの詳細
    all_ch_tables = []
    for insp in ok_files:
        for ch in insp.change_history_tables:
            ch["_file"] = insp.path
            all_ch_tables.append(ch)

    lines.append(f"  変更履歴テーブル検出数: {len(all_ch_tables)}個（{sum(1 for i in ok_files if i.change_history_tables)}ファイル）")
    lines.append(f"  検出条件: 1行目に「{'／'.join(sorted(CHANGE_HISTORY_KEYWORDS))}」のうち{CHANGE_HISTORY_MIN_MATCH}個以上を含む")
    lines.append("")

    if all_ch_tables:
        lines.append("  検出された変更履歴テーブル（全件）:")
        for ch in all_ch_tables:
            first_row_str = " | ".join(ch["first_row_texts"][:6])
            kw_str = ", ".join(ch["matched_keywords"])
            lines.append(f"    - {Path(ch['_file']).name} / 表{ch['table_index']}: {ch['rows']}行×{ch['cols']}列")
            lines.append(f"      1行目: [{first_row_str}]")
            lines.append(f"      一致キーワード: {kw_str}（{ch['match_count']}個）")
        lines.append("")

        # 1行目パターンの集約（変更履歴テーブルのみ）
        ch_patterns: Counter = Counter()
        for ch in all_ch_tables:
            key = " | ".join(ch["first_row_texts"][:6])
            ch_patterns[key] += 1

        if len(ch_patterns) > 1:
            lines.append("  変更履歴テーブルの1行目バリエーション:")
            for pattern, count in ch_patterns.most_common():
                lines.append(f"    [{count}回] {pattern}")
            lines.append("")

    # 検出漏れの可能性チェック: 変更履歴テーブルなしファイルの表1行目
    no_ch_files = [i for i in ok_files if not i.change_history_tables and i.tables]
    if no_ch_files:
        lines.append("  参考: 変更履歴テーブル未検出ファイルの表1行目パターン:")
        misc_patterns: Counter = Counter()
        for insp in no_ch_files:
            for t in insp.tables:
                if t.get("first_row_texts"):
                    key = " | ".join(t["first_row_texts"][:5])
                    misc_patterns[key] += 1
        for pattern, count in misc_patterns.most_common(10):
            lines.append(f"    [{count}回] {pattern}")
        lines.append("  → 上記に変更履歴らしきパターンがあれば検出条件の調整が必要")
    lines.append("")

    # --- 図形の詳細集計（一時コメントアウト） ---
    # lines.append("■ 図形の詳細分析")
    # lines.append("")
    #
    # all_shapes = []
    # for insp in ok_files:
    #     for s in insp.shapes:
    #         s["_file"] = insp.path
    #         all_shapes.append(s)
    #
    # if not all_shapes:
    #     lines.append("  図形: なし")
    # else:
    #     shape_types: Counter = Counter()
    #     for s in all_shapes:
    #         shape_types[s["shape_type"]] += 1
    #
    #     lines.append(f"  図形の総数: {len(all_shapes)}個")
    #     for st, count in shape_types.most_common():
    #         lines.append(f"    {st}: {count}個")
    #
    #     shapes_with_text = [s for s in all_shapes if s["has_text_content"]]
    #     lines.append(f"  テキストを含む図形: {len(shapes_with_text)}個")
    #
    #     if shapes_with_text:
    #         lines.append("  図形内テキストのサンプル:")
    #         for s in shapes_with_text[:10]:
    #             text = s.get("text", "")
    #             if text:
    #                 lines.append(f"    [{s['shape_type']}] 「{text[:60]}」")
    #
    #     groups = [s for s in all_shapes if s["shape_type"] == "group"]
    #     if groups:
    #         child_counts = [s["child_count"] for s in groups]
    #         lines.append(f"  グループ図形: {len(groups)}個（子要素数: 最小{min(child_counts)}, 最大{max(child_counts)}）")
    #         lines.append("  → ワークフロー図やフロー図の可能性が高い")
    # lines.append("")

    # --- 文書構造パターン（一時コメントアウト） ---
    # lines.append("■ 文書構造パターン")
    # lines.append("")
    #
    # for insp in ok_files:
    #     if not insp.element_order:
    #         continue
    #     type_sequence = [e["element_type"][0].upper() for e in insp.element_order[:30]]
    #     pseudo_count = sum(1 for e in insp.element_order if e["is_pseudo_heading"])
    #     lines.append(f"  {Path(insp.path).name}:")
    #     lines.append(f"    要素数: {len(insp.element_order)}、疑似見出し: {pseudo_count}個")
    #     lines.append(f"    先頭30要素: {''.join(type_sequence)} (P=段落, T=表)")
    #
    #     pt_pattern = 0
    #     for j in range(len(insp.element_order) - 1):
    #         if (
    #             insp.element_order[j]["element_type"] == "paragraph"
    #             and insp.element_order[j + 1]["element_type"] == "table"
    #         ):
    #             pt_pattern += 1
    #     if pt_pattern > 0:
    #         lines.append(f"    「段落→表」パターン: {pt_pattern}回")
    #     lines.append("")

    lines.append("=" * 60)
    lines.append("以上")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# メイン
# ---------------------------------------------------------------------------
def main() -> None:
    parser = argparse.ArgumentParser(
        description="docx 詳細調査: フォントパターン、表構造、図形、要素順序を調査する"
    )
    parser.add_argument("target", help=".docx ファイルまたはフォルダ")
    parser.add_argument(
        "--output", "-o", default="inspect_report",
        help="出力ファイル名（拡張子なし）。.json と .txt を出力 (デフォルト: inspect_report)",
    )
    args = parser.parse_args()

    target = Path(args.target)
    if target.is_file() and target.suffix.lower() == ".docx":
        docx_files = [target]
    elif target.is_dir():
        docx_files = sorted(target.rglob("*.docx"))
    else:
        print(f"エラー: {target} は .docx ファイルでもフォルダでもありません", file=sys.stderr)
        sys.exit(1)

    if not docx_files:
        print("対象の .docx ファイルが見つかりません", file=sys.stderr)
        sys.exit(1)

    print(f"調査開始: {len(docx_files)}ファイル")

    inspections: list[DocxInspection] = []
    for i, fpath in enumerate(docx_files, 1):
        print(f"  [{i}/{len(docx_files)}] {fpath.name}")
        insp = inspect_file(fpath)
        inspections.append(insp)

    print("調査完了")

    # JSON 出力
    json_path = Path(f"{args.output}.json")
    data = [asdict(insp) for insp in inspections]
    json_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"JSON レポート出力: {json_path}")

    # テキストレポート出力
    text_report = build_text_report(inspections)
    txt_path = Path(f"{args.output}.txt")
    txt_path.write_text(text_report, encoding="utf-8")
    print(f"テキストレポート出力: {txt_path}")

    # コンソール表示
    print("")
    print(text_report)


if __name__ == "__main__":
    main()
