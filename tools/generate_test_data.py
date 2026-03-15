"""テストデータ .docx 生成スクリプト

python-docx で作成可能なパターンのテストデータを生成する。

使い方:
  python tools/generate_test_data.py [--output-dir input/]
  # デフォルト出力先: input/word/

生成ファイル:
  1. many_tables.docx       — 表が多数あるドキュメント
  2. many_images.docx       — 画像が多数埋め込まれたドキュメント
  3. many_objects.docx      — InlineShape / XML 図形が多いドキュメント
  4. mixed_complex.docx     — 表・画像・図形・見出し混在の総合ドキュメント
  5. large_document.docx    — 大量コンテンツ（15MB 分割テスト用）
  6. oasys_style.docx       — Oasys/Win スタイル（見出しスタイルなし、フォントサイズ差のみ）
  7. change_history.docx    — 変更履歴テーブル付きドキュメント
"""

from __future__ import annotations

import argparse
import io
import struct
from functools import lru_cache
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Inches, Pt, RGBColor

# 名前空間定義
_NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "v": "urn:schemas-microsoft-com:vml",
    "o": "urn:schemas-microsoft-com:office:office",
}
_JP_FONT_CANDIDATES = [
    Path("C:/Windows/Fonts/msgothic.ttc"),
    Path("C:/Windows/Fonts/YuGothM.ttc"),
    Path("C:/Windows/Fonts/meiryo.ttc"),
    Path("C:/Windows/Fonts/MSMINCHO.TTC"),
]


def _vml_element(tag: str, **attribs) -> etree._Element:
    """VML/Office 名前空間を使って lxml 要素を直接作成する"""
    pfx, _, local = tag.partition(":")
    ns_uri = _NSMAP.get(pfx, _NSMAP["w"])
    el = etree.SubElement(etree.Element("dummy"), f"{{{ns_uri}}}{local}")
    # SubElement で作った要素を親から外す
    el.getparent().remove(el)
    for k, v in attribs.items():
        el.set(k, v)
    return el


@lru_cache(maxsize=16)
def _load_japanese_font(size: int):
    from PIL import ImageFont

    for font_path in _JP_FONT_CANDIDATES:
        if not font_path.exists():
            continue
        try:
            return ImageFont.truetype(str(font_path), size)
        except OSError:
            continue
    return ImageFont.load_default()


def _make_dummy_png(width: int = 100, height: int = 80, color: tuple = (70, 130, 180), *, noisy: bool = False) -> bytes:
    """最小限の PNG 画像をバイナリで生成する（Pillow 使用）

    noisy=True にするとランダムノイズを加え、PNG 圧縮が効きにくくなる（大容量テスト用）。
    """
    import random
    from PIL import Image

    img = Image.new("RGB", (width, height), color)
    if noisy:
        pixels = img.load()
        for y in range(height):
            for x in range(width):
                r, g, b = color
                pixels[x, y] = (
                    min(255, max(0, r + random.randint(-50, 50))),
                    min(255, max(0, g + random.randint(-50, 50))),
                    min(255, max(0, b + random.randint(-50, 50))),
                )
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _add_floating_textbox(paragraph, text: str, left_emu: int = 0, top_emu: int = 0):
    """浮動テキストボックスを VML pict 形式で追加する"""
    run = paragraph.add_run()

    pict = OxmlElement("w:pict")

    vshape = _vml_element("v:shape",
        style=f"position:absolute;left:{left_emu // 9525}pt;top:{top_emu // 9525}pt;width:150pt;height:40pt",
        type="#_x0000_t202",
        fillcolor="#dce6f1",
        strokecolor="#4472c4",
    )

    vtextbox = _vml_element("v:textbox")
    txbx_content = OxmlElement("w:txbxContent")
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    txbx_content.append(p)
    vtextbox.append(txbx_content)
    vshape.append(vtextbox)
    pict.append(vshape)

    run._element.append(pict)


# ─────────────────────────────────────────────
# 1. 表が多いドキュメント
# ─────────────────────────────────────────────
def generate_many_tables(output_dir: Path) -> Path:
    doc = Document()
    doc.add_heading("表が多数含まれるテストドキュメント", level=1)
    doc.add_paragraph("このドキュメントには 10 個以上の表が含まれる。")

    table_specs = [
        ("ユーザー管理テーブル", ["ID", "氏名", "メール", "権限"], [
            ["001", "田中太郎", "tanaka@example.com", "管理者"],
            ["002", "鈴木花子", "suzuki@example.com", "一般"],
            ["003", "佐藤次郎", "sato@example.com", "閲覧者"],
            ["004", "高橋美咲", "takahashi@example.com", "一般"],
        ]),
        ("エラーコード一覧", ["コード", "種別", "メッセージ", "対処法"], [
            ["E001", "入力", "必須項目が未入力です", "項目を入力してください"],
            ["E002", "入力", "形式が不正です", "正しい形式で入力してください"],
            ["E003", "DB", "接続に失敗しました", "DB サーバーの状態を確認"],
            ["E004", "DB", "タイムアウト", "ネットワーク状態を確認"],
            ["E005", "認証", "トークンが無効です", "再ログインしてください"],
            ["E006", "認証", "権限不足", "管理者に連絡"],
        ]),
        ("画面一覧", ["画面ID", "画面名", "概要"], [
            ["SC001", "ログイン画面", "認証を行う"],
            ["SC002", "ダッシュボード", "集計情報を表示"],
            ["SC003", "ユーザー一覧", "ユーザーの検索・表示"],
            ["SC004", "ユーザー詳細", "ユーザー情報の編集"],
            ["SC005", "帳票出力", "PDF/Excel 出力"],
        ]),
        ("API エンドポイント一覧", ["Method", "Path", "説明", "認証"], [
            ["GET", "/api/users", "ユーザー一覧取得", "要"],
            ["POST", "/api/users", "ユーザー作成", "要"],
            ["GET", "/api/users/{id}", "ユーザー詳細取得", "要"],
            ["PUT", "/api/users/{id}", "ユーザー更新", "要"],
            ["DELETE", "/api/users/{id}", "ユーザー削除", "要（管理者）"],
        ]),
        ("バッチ処理一覧", ["バッチID", "処理名", "スケジュール", "所要時間"], [
            ["B001", "日次集計", "毎日 02:00", "約30分"],
            ["B002", "月次締め", "毎月1日 01:00", "約2時間"],
            ["B003", "データ連携", "毎日 06:00/18:00", "約15分"],
        ]),
        ("テーブル定義: t_users", ["カラム名", "型", "NULL", "デフォルト", "説明"], [
            ["id", "BIGINT", "NO", "AUTO_INCREMENT", "主キー"],
            ["name", "VARCHAR(100)", "NO", "", "氏名"],
            ["email", "VARCHAR(255)", "NO", "", "メールアドレス"],
            ["role", "VARCHAR(20)", "NO", "'general'", "権限"],
            ["created_at", "DATETIME", "NO", "CURRENT_TIMESTAMP", "作成日時"],
            ["updated_at", "DATETIME", "YES", "NULL", "更新日時"],
            ["deleted_at", "DATETIME", "YES", "NULL", "論理削除日時"],
        ]),
        ("テーブル定義: t_logs", ["カラム名", "型", "NULL", "説明"], [
            ["id", "BIGINT", "NO", "主キー"],
            ["user_id", "BIGINT", "NO", "ユーザーID"],
            ["action", "VARCHAR(50)", "NO", "操作種別"],
            ["detail", "TEXT", "YES", "詳細"],
            ["created_at", "DATETIME", "NO", "記録日時"],
        ]),
        ("環境情報", ["環境", "サーバー", "DB", "URL"], [
            ["開発", "dev-app01", "dev-db01", "https://dev.example.com"],
            ["検証", "stg-app01", "stg-db01", "https://stg.example.com"],
            ["本番", "prd-app01/02", "prd-db01(M)/02(S)", "https://www.example.com"],
        ]),
        ("設定パラメータ", ["パラメータ", "デフォルト値", "説明"], [
            ["MAX_RETRY", "3", "リトライ回数上限"],
            ["TIMEOUT_SEC", "30", "タイムアウト秒数"],
            ["BATCH_SIZE", "1000", "バッチ処理単位"],
            ["LOG_LEVEL", "INFO", "ログレベル"],
            ["CACHE_TTL", "3600", "キャッシュ有効秒数"],
        ]),
        ("リリース計画", ["フェーズ", "内容", "開始", "終了", "担当"], [
            ["Phase1", "基本機能", "2025/04", "2025/06", "チームA"],
            ["Phase2", "帳票・連携", "2025/07", "2025/09", "チームB"],
            ["Phase3", "性能改善", "2025/10", "2025/12", "チームA"],
        ]),
    ]

    for caption, headers, data_rows in table_specs:
        doc.add_heading(caption, level=2)
        table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for r_idx, row_data in enumerate(data_rows):
            for c_idx, val in enumerate(row_data):
                table.rows[r_idx + 1].cells[c_idx].text = val
        doc.add_paragraph("")  # 表間の余白

    # 結合セルを含む表
    doc.add_heading("結合セルを含む表", level=2)
    t = doc.add_table(rows=5, cols=4)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "大分類"
    t.rows[0].cells[1].text = "中分類"
    t.rows[0].cells[2].text = "項目"
    t.rows[0].cells[3].text = "値"
    t.rows[1].cells[0].text = "入力系"
    t.rows[1].cells[1].text = "ファイル"
    t.rows[1].cells[2].text = "形式"
    t.rows[1].cells[3].text = "CSV"
    t.rows[2].cells[0].text = "入力系"
    t.rows[2].cells[1].text = "ファイル"
    t.rows[2].cells[2].text = "文字コード"
    t.rows[2].cells[3].text = "UTF-8"
    t.rows[3].cells[0].text = "出力系"
    t.rows[3].cells[1].text = "帳票"
    t.rows[3].cells[2].text = "形式"
    t.rows[3].cells[3].text = "PDF"
    t.rows[4].cells[0].text = "備考"
    t.rows[4].cells[0].merge(t.rows[4].cells[3])  # 横結合

    # 縦結合
    t.rows[1].cells[0].merge(t.rows[2].cells[0])  # "入力系" 縦結合

    path = output_dir / "many_tables.docx"
    doc.save(str(path))
    print(f"  生成: {path}")
    return path


# ─────────────────────────────────────────────
# 2. 画像が多いドキュメント
# ─────────────────────────────────────────────
def generate_many_images(output_dir: Path) -> Path:
    doc = Document()
    doc.add_heading("画像が多数含まれるテストドキュメント", level=1)
    doc.add_paragraph("このドキュメントには各セクションに画像が埋め込まれている。")

    image_sections = [
        ("システム構成図", (200, 150), (100, 149, 237)),
        ("ネットワーク構成図", (250, 180), (46, 139, 87)),
        ("画面遷移図", (200, 200), (220, 53, 69)),
        ("ER 図", (300, 200), (255, 165, 0)),
        ("クラス図", (250, 200), (147, 112, 219)),
        ("シーケンス図", (200, 300), (0, 128, 128)),
        ("デプロイ構成図", (250, 150), (128, 128, 0)),
        ("データフロー図", (300, 180), (199, 21, 133)),
    ]

    for title, (w, h), color in image_sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(f"以下に{title}を示す。")

        img_bytes = _make_dummy_png(w, h, color)
        img_stream = io.BytesIO(img_bytes)
        doc.add_picture(img_stream, width=Inches(4))

        doc.add_paragraph(f"図: {title}")
        doc.add_paragraph("")

    # 表の中に画像がある場合（セル内画像）
    doc.add_heading("アイコン一覧", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "アイコン"
    table.rows[0].cells[1].text = "名前"
    table.rows[0].cells[2].text = "用途"

    icons = [
        ("ユーザー", "ユーザー関連操作", (70, 130, 180)),
        ("設定", "システム設定", (46, 139, 87)),
        ("警告", "エラー・警告表示", (220, 53, 69)),
    ]
    for i, (name, usage, color) in enumerate(icons):
        img_bytes = _make_dummy_png(40, 40, color)
        img_stream = io.BytesIO(img_bytes)
        paragraph = table.rows[i + 1].cells[0].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(img_stream, width=Inches(0.5))
        table.rows[i + 1].cells[1].text = name
        table.rows[i + 1].cells[2].text = usage

    path = output_dir / "many_images.docx"
    doc.save(str(path))
    print(f"  生成: {path}")
    return path


# ─────────────────────────────────────────────
# 3. オブジェクト（図形）が多いドキュメント
# ─────────────────────────────────────────────
def generate_many_objects(output_dir: Path) -> Path:
    doc = Document()
    doc.add_heading("図形が多数含まれるテストドキュメント", level=1)
    doc.add_paragraph("InlineShape および浮動テキストボックスを含む。")

    # InlineShape（画像ベース）で図形を模擬
    doc.add_heading("フロー図", level=2)
    doc.add_paragraph("以下にフロー図を示す。")

    # フロー図の各ステップを画像で模擬
    flow_steps = ["開始", "入力チェック", "データ変換", "バリデーション", "DB登録", "結果通知", "終了"]
    for step in flow_steps:
        from PIL import Image, ImageDraw
        img = Image.new("RGB", (200, 50), (230, 230, 250))
        draw = ImageDraw.Draw(img)
        draw.rectangle([2, 2, 197, 47], outline=(100, 100, 180), width=2)
        font = _load_japanese_font(16)
        bbox = draw.textbbox((0, 0), step, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.text(((200 - tw) // 2, (50 - th) // 2), step, fill=(0, 0, 0), font=font)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        inline_shape = doc.add_picture(buf, width=Inches(1.8))
        # alt text を設定（抽出時にテキスト情報を保持するため）
        inline_shape._inline.docPr.set("descr", f"フロー図ステップ: {step}")

    doc.add_paragraph("↓ の矢印で接続")

    # 浮動テキストボックスを追加
    doc.add_heading("テキストボックス群", level=2)
    textboxes = [
        "処理A: データ取得",
        "処理B: 変換処理",
        "処理C: 出力処理",
        "判定: エラー有無",
        "処理D: エラーハンドリング",
        "処理E: ログ出力",
    ]
    for i, text in enumerate(textboxes):
        p = doc.add_paragraph()
        _add_floating_textbox(p, text, left_emu=i * 200000, top_emu=i * 100000)

    # VML 風の図形（古い Word 形式）
    doc.add_heading("VML 図形", level=2)
    doc.add_paragraph("以下は VML 形式の図形を含むセクション。")

    for label in ["入力", "処理", "出力"]:
        p = doc.add_paragraph()
        run = p.add_run()
        # pict > v:shape で VML 図形を追加
        pict = OxmlElement("w:pict")
        vshape = _vml_element("v:shape",
            style="width:100pt;height:40pt",
            type="#_x0000_t202",
        )
        vtextbox = _vml_element("v:textbox")
        txbx_content = OxmlElement("w:txbxContent")
        vp = OxmlElement("w:p")
        vr = OxmlElement("w:r")
        vt = OxmlElement("w:t")
        vt.text = label
        vr.append(vt)
        vp.append(vr)
        txbx_content.append(vp)
        vtextbox.append(txbx_content)
        vshape.append(vtextbox)
        pict.append(vshape)
        run._element.append(pict)

    path = output_dir / "many_objects.docx"
    doc.save(str(path))
    print(f"  生成: {path}")
    return path


# ─────────────────────────────────────────────
# 4. 総合ドキュメント（混在）
# ─────────────────────────────────────────────
def generate_mixed_complex(output_dir: Path) -> Path:
    doc = Document()
    doc.add_heading("機能仕様書（総合テスト）", level=1)
    doc.add_paragraph("本文書はシステムの機能仕様を定義する。")

    # 1章: 概要（見出し + 段落）
    doc.add_heading("1. 概要", level=2)
    doc.add_paragraph("本システムはデータ管理を行う Web アプリケーションである。")
    doc.add_paragraph("主要機能は以下の通り。")

    # 2章: 画像あり
    doc.add_heading("2. システム構成", level=2)
    doc.add_paragraph("以下にシステム構成図を示す。")
    img = _make_dummy_png(300, 200, (100, 149, 237))
    doc.add_picture(io.BytesIO(img), width=Inches(4))
    doc.add_paragraph("図2-1: システム構成図")

    # 3章: 表あり
    doc.add_heading("3. 画面一覧", level=2)
    t1 = doc.add_table(rows=4, cols=3)
    t1.style = "Table Grid"
    for i, h in enumerate(["画面ID", "画面名", "概要"]):
        t1.rows[0].cells[i].text = h
    t1.rows[1].cells[0].text = "SC001"
    t1.rows[1].cells[1].text = "ログイン"
    t1.rows[1].cells[2].text = "認証画面"
    t1.rows[2].cells[0].text = "SC002"
    t1.rows[2].cells[1].text = "一覧"
    t1.rows[2].cells[2].text = "データ一覧表示"
    t1.rows[3].cells[0].text = "SC003"
    t1.rows[3].cells[1].text = "詳細"
    t1.rows[3].cells[2].text = "データ詳細・編集"

    # 4章: フォントサイズ見出し + 表 + 画像
    doc.add_heading("4. 機能詳細", level=2)

    p = doc.add_paragraph()
    run = p.add_run("4.1 入力チェック機能")
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph("入力データの妥当性を検証する機能。")

    t2 = doc.add_table(rows=5, cols=3)
    t2.style = "Table Grid"
    for i, h in enumerate(["チェック項目", "条件", "エラーメッセージ"]):
        t2.rows[0].cells[i].text = h
    checks = [
        ["必須チェック", "値が空でないこと", "必須項目です"],
        ["桁数チェック", "最大100文字", "桁数超過です"],
        ["形式チェック", "半角英数字", "形式が不正です"],
        ["範囲チェック", "1〜9999", "範囲外の値です"],
    ]
    for r, row in enumerate(checks):
        for c, val in enumerate(row):
            t2.rows[r + 1].cells[c].text = val

    p2 = doc.add_paragraph()
    run2 = p2.add_run("4.2 データ出力機能")
    run2.font.size = Pt(14)
    run2.bold = True

    doc.add_paragraph("検証済みデータを CSV/PDF で出力する。")
    img2 = _make_dummy_png(250, 150, (46, 139, 87))
    doc.add_picture(io.BytesIO(img2), width=Inches(3.5))
    doc.add_paragraph("図4-2: 出力フロー")

    # 5章: 図形
    doc.add_heading("5. 処理フロー", level=2)
    p3 = doc.add_paragraph()
    _add_floating_textbox(p3, "開始", left_emu=0, top_emu=0)
    p4 = doc.add_paragraph()
    _add_floating_textbox(p4, "入力チェック", left_emu=200000, top_emu=0)
    p5 = doc.add_paragraph()
    _add_floating_textbox(p5, "終了", left_emu=400000, top_emu=0)

    # 変更履歴
    doc.add_paragraph("")
    doc.add_paragraph("変更履歴")
    ch = doc.add_table(rows=3, cols=4)
    ch.style = "Table Grid"
    ch.rows[0].cells[0].text = "ページ"
    ch.rows[0].cells[1].text = "種別"
    ch.rows[0].cells[2].text = "年月"
    ch.rows[0].cells[3].text = "記事"
    ch.rows[1].cells[0].text = "全"
    ch.rows[1].cells[1].text = "新規"
    ch.rows[1].cells[2].text = "2025/01"
    ch.rows[1].cells[3].text = "初版作成"
    ch.rows[2].cells[0].text = "5"
    ch.rows[2].cells[1].text = "修正"
    ch.rows[2].cells[2].text = "2025/06"
    ch.rows[2].cells[3].text = "入力チェック機能追加"

    path = output_dir / "mixed_complex.docx"
    doc.save(str(path))
    print(f"  生成: {path}")
    return path


# ─────────────────────────────────────────────
# 5. 大量コンテンツ（15MB 分割テスト用）
# ─────────────────────────────────────────────
def generate_large_document(output_dir: Path) -> Path:
    doc = Document()
    doc.add_heading("大規模仕様書（分割テスト用）", level=1)

    # 多数のセクション × 表 で容量を増やす
    for chapter in range(1, 21):
        doc.add_heading(f"第{chapter}章 機能{chapter}", level=2)
        doc.add_paragraph(f"本章では機能{chapter}の詳細仕様を記述する。" * 5)

        for section in range(1, 6):
            doc.add_heading(f"{chapter}.{section} サブ機能{section}", level=3)
            doc.add_paragraph(
                f"サブ機能{section}は、データの処理を行う。"
                f"入力されたデータに対して、バリデーション、変換、保存の一連の処理を実行する。"
                f"処理結果はログに記録され、エラー発生時にはリトライが行われる。" * 3
            )

            # 各サブセクションに表を追加
            t = doc.add_table(rows=6, cols=4)
            t.style = "Table Grid"
            for i, h in enumerate(["No", "項目", "条件", "備考"]):
                t.rows[0].cells[i].text = h
            for r in range(1, 6):
                t.rows[r].cells[0].text = str(r)
                t.rows[r].cells[1].text = f"チェック項目{r}"
                t.rows[r].cells[2].text = f"条件{r}の詳細説明文" * 2
                t.rows[r].cells[3].text = f"備考{r}"

        # 大きめのノイズ画像で容量を増やす
        img = _make_dummy_png(
            800, 600,
            ((chapter * 37) % 256, (chapter * 73) % 256, (chapter * 113) % 256),
            noisy=True,
        )
        doc.add_picture(io.BytesIO(img), width=Inches(5))

    path = output_dir / "large_document.docx"
    doc.save(str(path))
    size_mb = path.stat().st_size / (1024 * 1024)
    print(f"  生成: {path} ({size_mb:.1f} MB)")
    return path


# ─────────────────────────────────────────────
# 6. Oasys/Win スタイル（見出しスタイルなし）
# ─────────────────────────────────────────────
def generate_oasys_style(output_dir: Path) -> Path:
    doc = Document()

    # 見出しスタイルを使わず、フォントサイズのみで見出しを表現
    sections = [
        (16, "機能仕様書", True),
        (10.5, "本文書は機能Aの仕様を定義する。", False),
        (14, "1. 概要", True),
        (10.5, "本システムはデータ管理を目的とする。", False),
        (10.5, "対象ユーザーは社内の業務担当者である。", False),
        (14, "2. 機能一覧", True),
        (12, "2.1 入力機能", True),
        (10.5, "CSV ファイルを読み込み、データベースに登録する。", False),
        (10.5, "入力ファイルの文字コードは UTF-8 とする。", False),
        (12, "2.2 出力機能", True),
        (10.5, "データベースから抽出した結果を帳票として出力する。", False),
        (10.5, "出力形式は PDF および Excel とする。", False),
        (14, "3. 非機能要件", True),
        (12, "3.1 性能要件", True),
        (10.5, "レスポンスタイムは3秒以内とする。", False),
        (12, "3.2 可用性要件", True),
        (10.5, "稼働率 99.9% 以上を目標とする。", False),
        (14, "4. 制約事項", True),
        (10.5, "オフライン環境で動作すること。", False),
        (10.5, "Windows 10 以降で動作すること。", False),
    ]

    for size, text, is_bold in sections:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        if is_bold:
            run.bold = True

    # 表も追加
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("機能一覧表")
    run.font.size = Pt(12)
    run.bold = True

    t = doc.add_table(rows=4, cols=3)
    t.style = "Table Grid"
    for i, h in enumerate(["機能ID", "機能名", "優先度"]):
        t.rows[0].cells[i].text = h
    t.rows[1].cells[0].text = "F001"
    t.rows[1].cells[1].text = "入力チェック"
    t.rows[1].cells[2].text = "高"
    t.rows[2].cells[0].text = "F002"
    t.rows[2].cells[1].text = "データ変換"
    t.rows[2].cells[2].text = "高"
    t.rows[3].cells[0].text = "F003"
    t.rows[3].cells[1].text = "帳票出力"
    t.rows[3].cells[2].text = "中"

    path = output_dir / "oasys_style.docx"
    doc.save(str(path))
    print(f"  生成: {path}")
    return path


# ─────────────────────────────────────────────
# 7. 変更履歴テーブル付き
# ─────────────────────────────────────────────
def generate_change_history(output_dir: Path) -> Path:
    doc = Document()
    doc.add_heading("システム設計書", level=1)
    doc.add_paragraph("本文書はシステムの設計を記述する。")

    doc.add_heading("1. 設計方針", level=2)
    doc.add_paragraph("マイクロサービスアーキテクチャを採用する。")

    # 通常の表
    t = doc.add_table(rows=3, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "項目"
    t.rows[0].cells[1].text = "内容"
    t.rows[1].cells[0].text = "アーキテクチャ"
    t.rows[1].cells[1].text = "マイクロサービス"
    t.rows[2].cells[0].text = "DB"
    t.rows[2].cells[1].text = "PostgreSQL"

    doc.add_heading("2. コンポーネント設計", level=2)
    doc.add_paragraph("各コンポーネントの責務を以下に示す。")

    # 変更履歴テーブル（全角スペース入り — 実際の Oasys 出力パターン）
    doc.add_paragraph("")
    doc.add_paragraph("変更履歴")
    ch = doc.add_table(rows=5, cols=4)
    ch.style = "Table Grid"
    ch.rows[0].cells[0].text = "ページ"
    ch.rows[0].cells[1].text = "種別"
    ch.rows[0].cells[2].text = "年\u3000月"        # 全角スペース
    ch.rows[0].cells[3].text = "記\u3000\u3000事"  # 全角スペース2個
    ch.rows[1].cells[0].text = "全"
    ch.rows[1].cells[1].text = "新規"
    ch.rows[1].cells[2].text = "2024/10"
    ch.rows[1].cells[3].text = "初版作成"
    ch.rows[2].cells[0].text = "3"
    ch.rows[2].cells[1].text = "追加"
    ch.rows[2].cells[2].text = "2025/01"
    ch.rows[2].cells[3].text = "コンポーネント設計追加"
    ch.rows[3].cells[0].text = "5-8"
    ch.rows[3].cells[1].text = "修正"
    ch.rows[3].cells[2].text = "2025/03"
    ch.rows[3].cells[3].text = "API 設計見直し"
    ch.rows[4].cells[0].text = "10"
    ch.rows[4].cells[1].text = "削除"
    ch.rows[4].cells[2].text = "2025/06"
    ch.rows[4].cells[3].text = "旧バッチ処理削除"

    path = output_dir / "change_history.docx"
    doc.save(str(path))
    print(f"  生成: {path}")
    return path


# ─────────────────────────────────────────────
# 8. セル結合が複雑なドキュメント
# ─────────────────────────────────────────────
def generate_merged_cells(output_dir: Path) -> Path:
    doc = Document()
    doc.add_heading("セル結合テストドキュメント", level=1)
    doc.add_paragraph("縦結合・横結合・階層構造の結合パターンを含む。")

    # ── 表1: 縦結合（大分類→中分類の階層） ──
    doc.add_heading("1. 縦結合（大分類・中分類の階層）", level=2)
    t1 = doc.add_table(rows=8, cols=4)
    t1.style = "Table Grid"
    # ヘッダー
    for i, h in enumerate(["大分類", "中分類", "項目", "値"]):
        t1.rows[0].cells[i].text = h
    # 先に結合してからテキストを設定する（結合前にテキストがあると連結される）
    # 縦結合実行
    t1.rows[1].cells[0].merge(t1.rows[4].cells[0])  # 入力系: 4行結合
    t1.rows[1].cells[1].merge(t1.rows[2].cells[1])  # ファイル入力: 2行結合
    t1.rows[3].cells[1].merge(t1.rows[4].cells[1])  # DB入力: 2行結合
    t1.rows[5].cells[0].merge(t1.rows[7].cells[0])  # 出力系: 3行結合
    t1.rows[5].cells[1].merge(t1.rows[6].cells[1])  # 帳票: 2行結合
    # データ（結合済みセルの先頭行にのみテキスト設定）
    t1.rows[1].cells[0].text = "入力系"
    t1.rows[1].cells[1].text = "ファイル入力"
    t1.rows[1].cells[2].text = "形式"
    t1.rows[1].cells[3].text = "CSV"
    # row2: 大分類・中分類は結合済み（テキスト設定不要）
    t1.rows[2].cells[2].text = "文字コード"
    t1.rows[2].cells[3].text = "UTF-8"
    t1.rows[3].cells[1].text = "DB入力"
    t1.rows[3].cells[2].text = "接続先"
    t1.rows[3].cells[3].text = "PostgreSQL"
    # row4: 大分類・中分類は結合済み
    t1.rows[4].cells[2].text = "タイムアウト"
    t1.rows[4].cells[3].text = "30秒"
    t1.rows[5].cells[0].text = "出力系"
    t1.rows[5].cells[1].text = "帳票"
    t1.rows[5].cells[2].text = "形式"
    t1.rows[5].cells[3].text = "PDF"
    # row6: 大分類・中分類は結合済み
    t1.rows[6].cells[2].text = "用紙サイズ"
    t1.rows[6].cells[3].text = "A4"
    # row7: 大分類は結合済み
    t1.rows[7].cells[1].text = "ログ"
    t1.rows[7].cells[2].text = "出力先"
    t1.rows[7].cells[3].text = "/var/log/app"

    doc.add_paragraph("")

    # ── 表2: 横結合（セクションヘッダー） ──
    doc.add_heading("2. 横結合（セクションヘッダー）", level=2)
    t2 = doc.add_table(rows=7, cols=4)
    t2.style = "Table Grid"
    for i, h in enumerate(["項目", "設定値", "デフォルト", "備考"]):
        t2.rows[0].cells[i].text = h
    # セクションヘッダー行（横結合）
    t2.rows[1].cells[0].text = "■ 接続設定"
    t2.rows[1].cells[0].merge(t2.rows[1].cells[3])  # 4列結合
    t2.rows[2].cells[0].text = "ホスト名"
    t2.rows[2].cells[1].text = "db-server01"
    t2.rows[2].cells[2].text = "localhost"
    t2.rows[2].cells[3].text = ""
    t2.rows[3].cells[0].text = "ポート"
    t2.rows[3].cells[1].text = "5432"
    t2.rows[3].cells[2].text = "5432"
    t2.rows[3].cells[3].text = ""
    # 2番目のセクションヘッダー
    t2.rows[4].cells[0].text = "■ 認証設定"
    t2.rows[4].cells[0].merge(t2.rows[4].cells[3])  # 4列結合
    t2.rows[5].cells[0].text = "ユーザー名"
    t2.rows[5].cells[1].text = "app_user"
    t2.rows[5].cells[2].text = ""
    t2.rows[5].cells[3].text = "必須"
    t2.rows[6].cells[0].text = "認証方式"
    t2.rows[6].cells[1].text = "パスワード"
    t2.rows[6].cells[2].text = "パスワード"
    t2.rows[6].cells[3].text = ""

    doc.add_paragraph("")

    # ── 表3: 縦横結合の混在（マトリクス） ──
    doc.add_heading("3. 縦横結合の混在（テスト結果マトリクス）", level=2)
    t3 = doc.add_table(rows=7, cols=5)
    t3.style = "Table Grid"
    # 先に結合してからテキスト設定
    t3.rows[0].cells[1].merge(t3.rows[0].cells[2])  # 「テスト環境」2列結合
    t3.rows[0].cells[3].merge(t3.rows[0].cells[4])  # 「本番環境」2列結合
    t3.rows[0].cells[0].merge(t3.rows[1].cells[0])  # 「テスト項目」縦結合
    # テキスト設定
    t3.rows[0].cells[0].text = "テスト項目"
    t3.rows[0].cells[1].text = "テスト環境"
    t3.rows[0].cells[3].text = "本番環境"
    t3.rows[1].cells[1].text = "Windows"
    t3.rows[1].cells[2].text = "Linux"
    t3.rows[1].cells[3].text = "Windows"
    t3.rows[1].cells[4].text = "Linux"
    # データ行
    test_items = [
        ("機能テスト", "OK", "OK", "OK", "OK"),
        ("性能テスト", "OK", "NG", "未実施", "未実施"),
        ("負荷テスト", "OK", "OK", "NG", "OK"),
        ("セキュリティテスト", "OK", "OK", "OK", "OK"),
        ("回帰テスト", "OK", "OK", "OK", "NG"),
    ]
    for r, (item, *vals) in enumerate(test_items, start=2):
        t3.rows[r].cells[0].text = item
        for c, v in enumerate(vals):
            t3.rows[r].cells[c + 1].text = v

    doc.add_paragraph("")

    # ── 表4: 深い階層構造（3段の縦結合） ──
    doc.add_heading("4. 深い階層構造（3段の縦結合）", level=2)
    t4 = doc.add_table(rows=10, cols=5)
    t4.style = "Table Grid"
    for i, h in enumerate(["システム", "サブシステム", "モジュール", "機能", "状態"]):
        t4.rows[0].cells[i].text = h

    # 先に結合してからテキスト設定
    # Level 1: システム
    t4.rows[1].cells[0].merge(t4.rows[5].cells[0])  # 基幹系: 5行
    t4.rows[6].cells[0].merge(t4.rows[9].cells[0])  # 情報系: 4行
    # Level 2: サブシステム
    t4.rows[1].cells[1].merge(t4.rows[3].cells[1])  # 販売管理: 3行
    t4.rows[4].cells[1].merge(t4.rows[5].cells[1])  # 在庫管理: 2行
    t4.rows[6].cells[1].merge(t4.rows[8].cells[1])  # 帳票: 3行
    # Level 3: モジュール
    t4.rows[1].cells[2].merge(t4.rows[2].cells[2])  # 受注: 2行
    t4.rows[4].cells[2].merge(t4.rows[5].cells[2])  # 入庫: 2行
    t4.rows[6].cells[2].merge(t4.rows[7].cells[2])  # 月次: 2行

    # 結合済みセルの先頭行にのみテキスト設定
    t4.rows[1].cells[0].text = "基幹系"
    t4.rows[1].cells[1].text = "販売管理"
    t4.rows[1].cells[2].text = "受注"
    t4.rows[1].cells[3].text = "受注登録"
    t4.rows[1].cells[4].text = "稼働中"
    # row2: システム・サブシステム・モジュール結合済み
    t4.rows[2].cells[3].text = "受注変更"
    t4.rows[2].cells[4].text = "稼働中"
    # row3: システム・サブシステム結合済み
    t4.rows[3].cells[2].text = "出荷"
    t4.rows[3].cells[3].text = "出荷指示"
    t4.rows[3].cells[4].text = "稼働中"
    # row4: システム結合済み
    t4.rows[4].cells[1].text = "在庫管理"
    t4.rows[4].cells[2].text = "入庫"
    t4.rows[4].cells[3].text = "入庫登録"
    t4.rows[4].cells[4].text = "開発中"
    # row5: システム・サブシステム・モジュール結合済み
    t4.rows[5].cells[3].text = "入庫取消"
    t4.rows[5].cells[4].text = "開発中"
    # row6: 情報系の先頭
    t4.rows[6].cells[0].text = "情報系"
    t4.rows[6].cells[1].text = "帳票"
    t4.rows[6].cells[2].text = "月次"
    t4.rows[6].cells[3].text = "売上集計"
    t4.rows[6].cells[4].text = "稼働中"
    # row7: システム・サブシステム・モジュール結合済み
    t4.rows[7].cells[3].text = "在庫集計"
    t4.rows[7].cells[4].text = "テスト中"
    # row8: システム・サブシステム結合済み
    t4.rows[8].cells[2].text = "日次"
    t4.rows[8].cells[3].text = "日報出力"
    t4.rows[8].cells[4].text = "稼働中"
    # row9: システム結合済み
    t4.rows[9].cells[1].text = "ダッシュボード"
    t4.rows[9].cells[2].text = "KPI"
    t4.rows[9].cells[3].text = "KPI表示"
    t4.rows[9].cells[4].text = "企画中"

    doc.add_paragraph("")

    # ── 表5: 複合ヘッダー（横結合+縦結合のヘッダー） ──
    doc.add_heading("5. 複合ヘッダー（横結合+縦結合）", level=2)
    t5 = doc.add_table(rows=6, cols=6)
    t5.style = "Table Grid"
    # 先に結合してからテキスト設定
    t5.rows[0].cells[2].merge(t5.rows[0].cells[3])  # 入力: 2列結合
    t5.rows[0].cells[4].merge(t5.rows[0].cells[5])  # 出力: 2列結合
    t5.rows[0].cells[0].merge(t5.rows[1].cells[0])  # コード: 縦結合
    t5.rows[0].cells[1].merge(t5.rows[1].cells[1])  # 名称: 縦結合
    # テキスト設定
    t5.rows[0].cells[0].text = "コード"
    t5.rows[0].cells[1].text = "名称"
    t5.rows[0].cells[2].text = "入力"
    t5.rows[0].cells[4].text = "出力"
    t5.rows[1].cells[2].text = "ファイル"
    t5.rows[1].cells[3].text = "DB"
    t5.rows[1].cells[4].text = "画面"
    t5.rows[1].cells[5].text = "帳票"
    # データ行
    rows5 = [
        ("F001", "受注処理", "CSV", "SELECT/INSERT", "一覧画面", "受注伝票"),
        ("F002", "出荷処理", "-", "SELECT/UPDATE", "詳細画面", "出荷指示書"),
        ("F003", "在庫照会", "-", "SELECT", "照会画面", "-"),
        ("F004", "月次集計", "CSV", "SELECT", "-", "月次レポート"),
    ]
    for r, row in enumerate(rows5, start=2):
        for c, val in enumerate(row):
            t5.rows[r].cells[c].text = val

    doc.add_paragraph("")

    # ── 表6: 不規則な結合（合計行・備考行） ──
    doc.add_heading("6. 不規則な結合（合計行・備考行）", level=2)
    t6 = doc.add_table(rows=7, cols=4)
    t6.style = "Table Grid"
    for i, h in enumerate(["品目", "数量", "単価", "金額"]):
        t6.rows[0].cells[i].text = h
    t6.rows[1].cells[0].text = "サーバー"
    t6.rows[1].cells[1].text = "2"
    t6.rows[1].cells[2].text = "500,000"
    t6.rows[1].cells[3].text = "1,000,000"
    t6.rows[2].cells[0].text = "ストレージ"
    t6.rows[2].cells[1].text = "4"
    t6.rows[2].cells[2].text = "200,000"
    t6.rows[2].cells[3].text = "800,000"
    t6.rows[3].cells[0].text = "ネットワーク機器"
    t6.rows[3].cells[1].text = "1"
    t6.rows[3].cells[2].text = "300,000"
    t6.rows[3].cells[3].text = "300,000"
    # 小計行（左3列を結合）
    t6.rows[4].cells[0].text = "小計"
    t6.rows[4].cells[0].merge(t6.rows[4].cells[2])  # 3列結合
    t6.rows[4].cells[3].text = "2,100,000"
    # 消費税行
    t6.rows[5].cells[0].text = "消費税（10%）"
    t6.rows[5].cells[0].merge(t6.rows[5].cells[2])  # 3列結合
    t6.rows[5].cells[3].text = "210,000"
    # 合計行
    t6.rows[6].cells[0].text = "合計"
    t6.rows[6].cells[0].merge(t6.rows[6].cells[2])  # 3列結合
    t6.rows[6].cells[3].text = "2,310,000"

    path = output_dir / "merged_cells.docx"
    doc.save(str(path))
    print(f"  生成: {path}")
    return path


# ─────────────────────────────────────────────
# 9. オブジェクト + テキストボックス重ね置きワークフロー
# ─────────────────────────────────────────────
def generate_overlay_workflow(output_dir: Path) -> Path:
    """矩形オブジェクトの上にテキストボックスを重ねたワークフロー図。

    実際の業務文書でよくあるパターン:
      - 矩形 (v:rect) = テキストなしの図形オブジェクト
      - テキストボックス (v:shape type=_x0000_t202) = 矩形の上に重ねて配置
      - 矢印・線は別の v:line や v:shape で表現

    抽出時の課題:
      - 矩形とテキストボックスは別要素なので対応関係が分からない
      - 位置情報 (style の left/top) から近接判定するしかない
      - テキストなし矩形だけ抽出しても意味がない
    """
    doc = Document()
    doc.add_heading("ワークフロー図（オブジェクト重ね置き）", level=1)
    doc.add_paragraph("矩形オブジェクトの上にテキストボックスを重ねたパターン。")

    # ── ワークフロー1: 承認フロー ──
    doc.add_heading("1. 承認フロー", level=2)
    doc.add_paragraph("以下に承認フローを示す。")

    # 各ステップ: 矩形（テキストなし）＋ テキストボックス（重ね置き）
    workflow_steps = [
        ("申請者\n申請書作成", 0, 0),
        ("上長\n内容確認", 200, 0),
        ("部長\n承認判断", 400, 0),
        ("経理\n処理実行", 600, 0),
        ("完了", 800, 0),
    ]

    for text, left, top in workflow_steps:
        p = doc.add_paragraph()
        run = p.add_run()

        # 1. 矩形オブジェクト（テキストなし）— v:rect
        pict_rect = OxmlElement("w:pict")
        vrect = _vml_element("v:rect",
            style=f"position:absolute;left:{left}pt;top:{top}pt;width:120pt;height:50pt",
            fillcolor="#d9e2f3",
            strokecolor="#4472c4",
            strokeweight="1.5pt",
        )
        pict_rect.append(vrect)
        run._element.append(pict_rect)

        # 2. テキストボックス（矩形の上に重ね置き）— v:shape
        p2 = doc.add_paragraph()
        _add_floating_textbox(p2, text, left_emu=left * 9525, top_emu=top * 9525)

    # 矢印を段落テキストで表現（VML の v:line は複雑なので簡略化）
    doc.add_paragraph("→ → → → →")
    doc.add_paragraph("")

    # ── ワークフロー2: エラーハンドリングフロー（分岐あり） ──
    doc.add_heading("2. エラーハンドリングフロー", level=2)
    doc.add_paragraph("以下にエラーハンドリングフローを示す。")

    error_flow = [
        # (text, left_pt, top_pt) — 分岐を含む2段構成
        ("データ受信", 0, 0),
        ("形式チェック", 200, 0),
        ("判定:\nエラー有無", 400, 0),
        # 正常系
        ("DB登録", 600, 0),
        ("完了通知", 800, 0),
        # エラー系（下段）
        ("エラーログ\n出力", 400, 80),
        ("管理者通知", 600, 80),
        ("リトライ待ち", 800, 80),
    ]

    for text, left, top in error_flow:
        p = doc.add_paragraph()
        run = p.add_run()

        # 判定ボックスはひし形風に色を変える
        is_decision = "判定" in text
        fill = "#fce4ec" if is_decision else "#e8f5e9"
        stroke = "#c62828" if is_decision else "#2e7d32"

        pict_rect = OxmlElement("w:pict")
        vrect = _vml_element("v:rect",
            style=f"position:absolute;left:{left}pt;top:{top}pt;width:130pt;height:50pt",
            fillcolor=fill,
            strokecolor=stroke,
            strokeweight="1.5pt",
        )
        pict_rect.append(vrect)
        run._element.append(pict_rect)

        p2 = doc.add_paragraph()
        _add_floating_textbox(p2, text, left_emu=left * 9525, top_emu=top * 9525)

    doc.add_paragraph("正常系: → → → →")
    doc.add_paragraph("エラー系: ↓ → → →")
    doc.add_paragraph("")

    # ── ワークフロー3: テキストボックスのみ（矩形なし） ──
    doc.add_heading("3. テキストボックスのみのフロー", level=2)
    doc.add_paragraph("矩形なし、テキストボックスだけで構成されたフロー。")

    textbox_only = [
        "受付登録",
        "データ検証",
        "変換処理",
        "結果出力",
    ]
    for i, text in enumerate(textbox_only):
        p = doc.add_paragraph()
        _add_floating_textbox(p, text, left_emu=i * 200 * 9525, top_emu=0)

    doc.add_paragraph("")

    # ── 説明文 ──
    doc.add_heading("4. 抽出上の課題", level=2)
    doc.add_paragraph(
        "上記のワークフロー図では、矩形オブジェクト（テキストなし）と"
        "テキストボックス（テキストあり）が同じ位置に重ねて配置されている。"
    )
    doc.add_paragraph(
        "抽出器は矩形とテキストボックスを別々の要素として検出するため、"
        "両者の対応関係を位置情報から推定する必要がある。"
    )
    doc.add_paragraph(
        "現状のパイプラインでは、テキストボックスのテキストは抽出できるが、"
        "テキストなし矩形はプレースホルダーとして出力される。"
    )

    path = output_dir / "overlay_workflow.docx"
    doc.save(str(path))
    print(f"  生成: {path}")
    return path


def _resolve_output_dir(base_dir: Path, leaf_name: str) -> Path:
    if base_dir.name.lower() == leaf_name.lower():
        return base_dir
    return base_dir / leaf_name


def main():
    parser = argparse.ArgumentParser(description="テストデータ .docx 生成")
    parser.add_argument(
        "--output-dir", "-o",
        type=Path, default=Path("input"),
        help="出力先ベースディレクトリ (default: input/ -> input/word/)",
    )
    args = parser.parse_args()

    output_dir = _resolve_output_dir(args.output_dir, "word")
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"テストデータ生成先: {output_dir}/")
    print()

    generate_many_tables(output_dir)
    generate_many_images(output_dir)
    generate_many_objects(output_dir)
    generate_mixed_complex(output_dir)
    generate_large_document(output_dir)
    generate_oasys_style(output_dir)
    generate_change_history(output_dir)
    generate_merged_cells(output_dir)
    generate_overlay_workflow(output_dir)

    print()
    print("完了: 9 ファイル生成しました。")


if __name__ == "__main__":
    main()
