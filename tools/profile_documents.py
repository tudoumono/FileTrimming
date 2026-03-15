"""文書プロファイリングスクリプト

指定フォルダ配下の全ファイルをスキャンし、拡張子ごとに取得可能な
構造情報を抽出して JSON レポートを出力する。

対象:
  - .docx: 見出し数・階層、表の数・行列数・結合セル、図形数、先頭見出し、文字数
  - .xlsx: シート数、各シートの行列数・空率
  - .doc / .rtf / .xls: ファイルサイズのみ（COM 変換なし）
  - BAGLES / PDF / その他: ファイルサイズのみ

使い方:
  python tools/profile_documents.py <input_dir> [--output profile_report.json]
"""

from __future__ import annotations

import argparse
import json
import math
import os
import sys
from collections import Counter
from dataclasses import asdict, dataclass, field
from pathlib import Path


# ---------------------------------------------------------------------------
# データクラス
# ---------------------------------------------------------------------------
@dataclass
class TableInfo:
    rows: int
    cols: int
    has_merged_cells: bool


@dataclass
class SheetInfo:
    name: str
    rows: int
    cols: int
    non_empty_cells: int
    total_cells: int
    empty_ratio: float


@dataclass
class FileProfile:
    path: str  # input_dir からの相対パス
    extension: str
    size_bytes: int
    category: str  # Word / Excel / BAGLES / PDF / PowerPoint / Text / Other

    # .docx 詳細
    heading_count: int | None = None
    heading_max_level: int | None = None
    headings: list[dict] | None = None  # [{level, text}, ...]
    first_heading: str | None = None
    paragraph_count: int | None = None
    char_count: int | None = None
    table_count: int | None = None
    tables: list[dict] | None = None  # [TableInfo as dict, ...]
    shape_count: int | None = None
    inline_shape_count: int | None = None

    # .xlsx 詳細
    sheet_count: int | None = None
    sheets: list[dict] | None = None  # [SheetInfo as dict, ...]
    hidden_sheet_count: int | None = None
    protected_sheet_count: int | None = None
    total_merged_cells: int | None = None
    total_tables: int | None = None
    total_images: int | None = None
    total_formulas: int | None = None
    total_comments: int | None = None
    has_named_ranges: bool | None = None
    named_range_count: int | None = None
    has_data_validation: bool | None = None
    has_conditional_formatting: bool | None = None
    has_print_settings: bool | None = None
    has_outline: bool | None = None
    has_auto_filter: bool | None = None
    estimated_layout_type: str | None = None

    # エラー
    error: str | None = None


# ---------------------------------------------------------------------------
# 拡張子 → カテゴリ マッピング
# ---------------------------------------------------------------------------
EXTENSION_CATEGORY: dict[str, str] = {
    ".docx": "Word",
    ".doc": "Word",
    ".rtf": "RTF",
    ".xlsx": "Excel",
    ".xls": "Excel",
    ".xlsm": "Excel",
    ".pptx": "PowerPoint",
    ".ppt": "PowerPoint",
    ".pdf": "PDF",
    ".bik": "BAGLES",
    ".bca": "BAGLES",
    ".bci": "BAGLES",
    ".bpg": "BAGLES",
    ".txt": "Text",
    ".csv": "Text",
    ".tsv": "Text",
    ".md": "Text",
    ".log": "Text",
}


def get_category(ext: str) -> str:
    return EXTENSION_CATEGORY.get(ext.lower(), "Other")


# ---------------------------------------------------------------------------
# .docx プロファイリング
# ---------------------------------------------------------------------------
def profile_docx(filepath: Path) -> dict:
    try:
        from docx import Document  # python-docx
    except ImportError:
        return {"error": "python-docx がインストールされていません (pip install python-docx)"}

    try:
        doc = Document(str(filepath))
    except Exception as e:
        return {"error": f"docx 読み込み失敗: {e}"}

    # 見出し
    headings: list[dict] = []
    char_count = 0
    paragraph_count = 0

    for para in doc.paragraphs:
        paragraph_count += 1
        char_count += len(para.text)
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                level = 0
            headings.append({"level": level, "text": para.text[:100]})

    heading_levels = [h["level"] for h in headings if h["level"] > 0]

    # 表
    tables: list[dict] = []
    for tbl in doc.tables:
        rows = len(tbl.rows)
        cols = len(tbl.columns)
        # 結合セルの検出: セルの実数とグリッドサイズの比較
        actual_cells = sum(1 for row in tbl.rows for _ in row.cells)
        grid_cells = rows * cols
        has_merged = actual_cells != grid_cells
        tables.append(asdict(TableInfo(rows=rows, cols=cols, has_merged_cells=has_merged)))

    # 図形
    inline_shape_count = len(doc.inline_shapes)

    # Shape (浮動図形) は python-docx の公開 API では直接取れないので
    # document.xml 内の wps:wsp / mc:AlternateContent を数える
    shape_count = 0
    try:
        from lxml import etree

        body = doc.element.body
        namespaces = {
            "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
            "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
            "v": "urn:schemas-microsoft-com:vml",
        }
        shape_count += len(body.findall(".//wps:wsp", namespaces))
        shape_count += len(body.findall(".//v:shape", namespaces))
    except Exception:
        shape_count = None  # lxml がなければスキップ

    return {
        "heading_count": len(headings),
        "heading_max_level": max(heading_levels) if heading_levels else 0,
        "headings": headings,
        "first_heading": headings[0]["text"] if headings else None,
        "paragraph_count": paragraph_count,
        "char_count": char_count,
        "table_count": len(tables),
        "tables": tables,
        "shape_count": shape_count,
        "inline_shape_count": inline_shape_count,
    }


# ---------------------------------------------------------------------------
# .xlsx プロファイリング
# ---------------------------------------------------------------------------
XLSX_FORMULA_SAMPLE_ROW_LIMIT = 1000
XLSX_LARGE_SHEET_ROW_THRESHOLD = 100
XLSX_LARGE_SHEET_COL_THRESHOLD = 20
XLSX_FORM_GRID_STDDEV_THRESHOLD = 0.5


def _xlsx_outline_level(dimension) -> int:
    return int(getattr(dimension, "outlineLevel", getattr(dimension, "outline_level", 0)) or 0)


def _xlsx_stddev(values: list[float]) -> float | None:
    if not values:
        return None
    avg = sum(values) / len(values)
    variance = sum((value - avg) ** 2 for value in values) / len(values)
    return math.sqrt(variance)


def _estimate_xlsx_layout_type(candidates: set[str]) -> str:
    if len(candidates) > 1:
        return "mixed"
    if len(candidates) == 1:
        return next(iter(candidates))
    return "unknown"


def profile_xlsx(filepath: Path) -> dict:
    try:
        from openpyxl import load_workbook
        from openpyxl.cell.cell import MergedCell
    except ImportError:
        return {"error": "openpyxl がインストールされていません (pip install openpyxl)"}

    try:
        wb = load_workbook(str(filepath), read_only=False, data_only=False)
    except Exception as e:
        return {"error": f"xlsx 読み込み失敗: {e}"}

    sheets: list[dict] = []
    hidden_sheet_count = 0
    protected_sheet_count = 0
    total_merged_cells = 0
    total_tables = 0
    total_images = 0
    total_formulas = 0
    total_comments = 0
    has_data_validation = False
    has_conditional_formatting = False
    has_print_settings = False
    has_outline = False
    has_auto_filter = False
    layout_candidates: set[str] = set()

    try:
        named_range_count = len(list(wb.defined_names.items()))

        for ws in wb.worksheets:
            max_row = ws.max_row or 0
            max_col = ws.max_column or 0
            total = max_row * max_col

            concrete_cells = [cell for cell in getattr(ws, "_cells", {}).values() if not isinstance(cell, MergedCell)]
            non_empty = sum(1 for cell in concrete_cells if cell.value is not None)
            empty_ratio = round(1 - (non_empty / total), 4) if total > 0 else 0.0
            sheets.append(
                asdict(
                    SheetInfo(
                        name=ws.title,
                        rows=max_row,
                        cols=max_col,
                        non_empty_cells=non_empty,
                        total_cells=total,
                        empty_ratio=empty_ratio,
                    )
                )
            )

            if ws.sheet_state != "visible":
                hidden_sheet_count += 1
            if ws.protection.sheet:
                protected_sheet_count += 1

            total_merged_cells += len(ws.merged_cells.ranges)
            total_tables += len(ws.tables)
            total_images += len(getattr(ws, "_images", []))
            total_comments += sum(1 for cell in concrete_cells if cell.comment is not None)

            sampled_formulas = sum(
                1
                for cell in concrete_cells
                if cell.row <= XLSX_FORMULA_SAMPLE_ROW_LIMIT
                and isinstance(cell.value, str)
                and cell.value.startswith("=")
            )
            total_formulas += sampled_formulas

            dv_count = len(getattr(ws.data_validations, "dataValidation", []))
            if dv_count > 0:
                has_data_validation = True

            cf_rule_count = 0
            if hasattr(ws.conditional_formatting, "_cf_rules"):
                cf_rule_count = sum(len(rules) for rules in ws.conditional_formatting._cf_rules.values())
            else:
                try:
                    cf_rule_count = len(ws.conditional_formatting)
                except Exception:
                    cf_rule_count = 0
            if cf_rule_count > 0:
                has_conditional_formatting = True

            if ws.print_area or ws.print_title_rows or ws.print_title_cols:
                has_print_settings = True
                layout_candidates.add("report_print")

            if ws.auto_filter and ws.auto_filter.ref:
                has_auto_filter = True

            row_outline = any(_xlsx_outline_level(dim) > 0 for dim in ws.row_dimensions.values())
            col_outline = any(_xlsx_outline_level(dim) > 0 for dim in ws.column_dimensions.values())
            if row_outline or col_outline:
                has_outline = True

            widths = [float(dim.width) for dim in ws.column_dimensions.values() if dim.width is not None]
            width_stddev = _xlsx_stddev(widths)
            grid_lines_visible = ws.sheet_view.showGridLines
            if grid_lines_visible is None:
                grid_lines_visible = True
            if width_stddev is not None and width_stddev <= XLSX_FORM_GRID_STDDEV_THRESHOLD and not grid_lines_visible:
                layout_candidates.add("form_grid")

            if len(ws.tables) > 0 or max_row > XLSX_LARGE_SHEET_ROW_THRESHOLD:
                layout_candidates.add("data_table")

        return {
            "sheet_count": len(sheets),
            "sheets": sheets,
            "hidden_sheet_count": hidden_sheet_count,
            "protected_sheet_count": protected_sheet_count,
            "total_merged_cells": total_merged_cells,
            "total_tables": total_tables,
            "total_images": total_images,
            "total_formulas": total_formulas,
            "total_comments": total_comments,
            "has_named_ranges": named_range_count > 0,
            "named_range_count": named_range_count,
            "has_data_validation": has_data_validation,
            "has_conditional_formatting": has_conditional_formatting,
            "has_print_settings": has_print_settings,
            "has_outline": has_outline,
            "has_auto_filter": has_auto_filter,
            "estimated_layout_type": _estimate_xlsx_layout_type(layout_candidates),
        }
    finally:
        wb.close()


# ---------------------------------------------------------------------------
# メイン処理
# ---------------------------------------------------------------------------
def scan_directory(input_dir: Path) -> list[FileProfile]:
    profiles: list[FileProfile] = []

    for root, _dirs, files in os.walk(input_dir):
        for fname in files:
            fpath = Path(root) / fname
            ext = fpath.suffix.lower()
            category = get_category(ext)
            rel_path = str(fpath.relative_to(input_dir))

            profile = FileProfile(
                path=rel_path,
                extension=ext,
                size_bytes=fpath.stat().st_size,
                category=category,
            )

            if ext == ".docx":
                details = profile_docx(fpath)
                if "error" in details:
                    profile.error = details["error"]
                else:
                    for k, v in details.items():
                        setattr(profile, k, v)

            elif ext in (".xlsx", ".xlsm"):
                details = profile_xlsx(fpath)
                if "error" in details:
                    profile.error = details["error"]
                else:
                    for k, v in details.items():
                        setattr(profile, k, v)

            profiles.append(profile)

    return profiles


def build_summary(profiles: list[FileProfile]) -> dict:
    """カテゴリ別・拡張子別の集計サマリーを生成する。"""
    by_category: dict[str, list] = {}
    by_extension: dict[str, list] = {}

    for p in profiles:
        by_category.setdefault(p.category, []).append(p)
        by_extension.setdefault(p.extension, []).append(p)

    summary: dict = {
        "total_files": len(profiles),
        "total_size_bytes": sum(p.size_bytes for p in profiles),
        "by_category": {},
        "by_extension": {},
    }

    for cat, items in sorted(by_category.items()):
        sizes = [i.size_bytes for i in items]
        summary["by_category"][cat] = {
            "count": len(items),
            "total_size_bytes": sum(sizes),
            "max_size_bytes": max(sizes),
            "min_size_bytes": min(sizes),
        }

    for ext, items in sorted(by_extension.items()):
        sizes = [i.size_bytes for i in items]
        summary["by_extension"][ext] = {
            "count": len(items),
            "total_size_bytes": sum(sizes),
            "max_size_bytes": max(sizes),
        }

    # 15MB 超のファイル一覧
    over_15mb = [p.path for p in profiles if p.size_bytes > 15 * 1024 * 1024]
    summary["files_over_15mb"] = over_15mb
    summary["files_over_15mb_count"] = len(over_15mb)

    # エラーが出たファイル一覧
    errors = [{"path": p.path, "error": p.error} for p in profiles if p.error]
    summary["errors"] = errors
    summary["error_count"] = len(errors)

    return summary


def format_size(size_bytes: int) -> str:
    """バイト数を読みやすい単位に変換する。"""
    if size_bytes >= 1024 * 1024:
        return f"{size_bytes / 1024 / 1024:.1f}MB"
    elif size_bytes >= 1024:
        return f"{size_bytes / 1024:.0f}KB"
    else:
        return f"{size_bytes}バイト"


def build_text_report(profiles: list[FileProfile], summary: dict) -> str:
    """口頭で説明しやすい日本語テキストレポートを生成する。"""
    lines: list[str] = []

    # --- 全体概要 ---
    lines.append("=" * 60)
    lines.append("文書プロファイリング結果")
    lines.append("=" * 60)
    lines.append("")
    total_size = format_size(summary["total_size_bytes"])
    lines.append(f"全体: {summary['total_files']}ファイル、合計{total_size}")
    lines.append("")

    # --- カテゴリ別概要 ---
    lines.append("■ カテゴリ別の内訳")
    lines.append("")
    for cat, info in summary["by_category"].items():
        max_s = format_size(info["max_size_bytes"])
        min_s = format_size(info["min_size_bytes"])
        lines.append(f"  {cat}: {info['count']}件 (最大{max_s}、最小{min_s})")
    lines.append("")

    # --- 15MB超 ---
    if summary["files_over_15mb"]:
        lines.append(f"■ 15MB超のファイル: {summary['files_over_15mb_count']}件")
        lines.append("  → Dify UI の上限を超えるため物理分割が必要")
        for f in summary["files_over_15mb"]:
            p = next((x for x in profiles if x.path == f), None)
            if p:
                lines.append(f"  - {f} ({format_size(p.size_bytes)})")
        lines.append("")
    else:
        lines.append("■ 15MB超のファイル: なし（全ファイル Dify UI の上限内）")
        lines.append("")

    # --- docx 詳細 ---
    docx_files = [p for p in profiles if p.extension == ".docx" and p.error is None]
    if docx_files:
        lines.append("■ docx ファイルの構造分析")
        lines.append("")

        # 見出し
        has_headings = [p for p in docx_files if p.heading_count and p.heading_count > 0]
        no_headings = [p for p in docx_files if not p.heading_count or p.heading_count == 0]
        lines.append(f"  見出しあり: {len(has_headings)}件 / 見出しなし: {len(no_headings)}件")
        if has_headings:
            avg_headings = sum(p.heading_count for p in has_headings) / len(has_headings)
            max_headings = max(p.heading_count for p in has_headings)
            max_level = max((p.heading_max_level or 0) for p in has_headings)
            lines.append(f"  見出し数: 平均{avg_headings:.0f}個、最大{max_headings}個、最深レベル{max_level}")
        lines.append("")

        # 表
        has_tables = [p for p in docx_files if p.table_count and p.table_count > 0]
        lines.append(f"  表を含むファイル: {len(has_tables)}件 / {len(docx_files)}件中")
        if has_tables:
            total_tables = sum(p.table_count for p in has_tables)
            avg_tables = total_tables / len(has_tables)
            max_tables = max(p.table_count for p in has_tables)
            lines.append(f"  表の数: 合計{total_tables}個、平均{avg_tables:.0f}個/ファイル、最大{max_tables}個")

            # 結合セル
            merged_count = 0
            for p in has_tables:
                if p.tables:
                    for t in p.tables:
                        if t.get("has_merged_cells"):
                            merged_count += 1
            if merged_count > 0:
                lines.append(f"  結合セルを含む表: {merged_count}個 → 変換時に注意が必要")
            else:
                lines.append("  結合セルを含む表: なし")

            # 大きな表
            large_tables = []
            for p in has_tables:
                if p.tables:
                    for t in p.tables:
                        if t["rows"] >= 20 or t["cols"] >= 10:
                            large_tables.append({"file": p.path, **t})
            if large_tables:
                lines.append(f"  大きな表（20行以上 or 10列以上）: {len(large_tables)}個")
                for lt in large_tables[:5]:
                    lines.append(f"    - {lt['file']}: {lt['rows']}行×{lt['cols']}列")
                if len(large_tables) > 5:
                    lines.append(f"    - 他{len(large_tables) - 5}個")
        lines.append("")

        # 図形
        has_shapes = [
            p for p in docx_files
            if (p.shape_count and p.shape_count > 0) or (p.inline_shape_count and p.inline_shape_count > 0)
        ]
        lines.append(f"  図形・画像を含むファイル: {len(has_shapes)}件 / {len(docx_files)}件中")
        if has_shapes:
            total_shapes = sum((p.shape_count or 0) for p in has_shapes)
            total_inline = sum((p.inline_shape_count or 0) for p in has_shapes)
            lines.append(f"  浮動図形: 合計{total_shapes}個、インライン画像: 合計{total_inline}個")
        lines.append("")

        # 文字数
        char_counts = [p.char_count for p in docx_files if p.char_count is not None]
        if char_counts:
            avg_chars = sum(char_counts) / len(char_counts)
            max_chars = max(char_counts)
            min_chars = min(char_counts)
            lines.append(f"  文字数: 平均{avg_chars:.0f}字、最大{max_chars}字、最小{min_chars}字")
            lines.append("")

        # 先頭見出しのパターン（doc_role 推定用）
        first_headings: dict[str, int] = {}
        for p in docx_files:
            fh = p.first_heading or "(見出しなし)"
            # 長い見出しは先頭30字で丸める
            key = fh[:30] if len(fh) > 30 else fh
            first_headings[key] = first_headings.get(key, 0) + 1
        if first_headings:
            lines.append("  先頭見出しの出現パターン（文書種別の推定材料）:")
            for heading, count in sorted(first_headings.items(), key=lambda x: -x[1]):
                lines.append(f"    「{heading}」: {count}件")
            lines.append("")

    # --- xlsx 詳細 ---
    xlsx_files = [p for p in profiles if p.extension in (".xlsx", ".xlsm") and p.error is None]
    if xlsx_files:
        lines.append("■ xlsx ファイルの構造分析")
        lines.append("")
        total_sheets = sum((p.sheet_count or 0) for p in xlsx_files)
        avg_sheets = total_sheets / len(xlsx_files) if xlsx_files else 0
        lines.append(f"  ファイル数: {len(xlsx_files)}件、シート数合計: {total_sheets}、平均{avg_sheets:.1f}シート/ファイル")

        # 大きなシート
        large_sheets = []
        for p in xlsx_files:
            if p.sheets:
                for s in p.sheets:
                    if s["rows"] >= XLSX_LARGE_SHEET_ROW_THRESHOLD or s["cols"] >= XLSX_LARGE_SHEET_COL_THRESHOLD:
                        large_sheets.append({"file": p.path, **s})
        if large_sheets:
            lines.append(f"  大きなシート（100行以上 or 20列以上）: {len(large_sheets)}件")
            for ls in large_sheets[:5]:
                lines.append(f"    - {ls['file']} / {ls['name']}: {ls['rows']}行×{ls['cols']}列 (空率{ls['empty_ratio']:.0%})")
            if len(large_sheets) > 5:
                lines.append(f"    - 他{len(large_sheets) - 5}件")

        hidden_sheet_files = [p for p in xlsx_files if (p.hidden_sheet_count or 0) > 0]
        protected_sheet_files = [p for p in xlsx_files if (p.protected_sheet_count or 0) > 0]
        merged_files = [p for p in xlsx_files if (p.total_merged_cells or 0) > 0]
        table_files = [p for p in xlsx_files if (p.total_tables or 0) > 0]
        image_files = [p for p in xlsx_files if (p.total_images or 0) > 0]
        formula_files = [p for p in xlsx_files if (p.total_formulas or 0) > 0]
        comment_files = [p for p in xlsx_files if (p.total_comments or 0) > 0]
        named_range_files = [p for p in xlsx_files if p.has_named_ranges]
        dv_files = [p for p in xlsx_files if p.has_data_validation]
        cf_files = [p for p in xlsx_files if p.has_conditional_formatting]
        print_setting_files = [p for p in xlsx_files if p.has_print_settings]
        outline_files = [p for p in xlsx_files if p.has_outline]
        filter_files = [p for p in xlsx_files if p.has_auto_filter]

        lines.append(f"  非表示シートを含むファイル: {len(hidden_sheet_files)}件")
        lines.append(f"  シート保護を含むファイル: {len(protected_sheet_files)}件")
        lines.append(
            f"  結合セルを含むファイル: {len(merged_files)}件 "
            f"(合計{sum((p.total_merged_cells or 0) for p in merged_files)}件)"
        )
        lines.append(f"  Excel Table を含むファイル: {len(table_files)}件")
        lines.append(f"  画像を含むファイル: {len(image_files)}件")
        lines.append(f"  数式を含むファイル: {len(formula_files)}件")
        lines.append(f"  コメントを含むファイル: {len(comment_files)}件")
        lines.append(f"  名前定義を含むファイル: {len(named_range_files)}件")
        lines.append(f"  入力規則を含むファイル: {len(dv_files)}件")
        lines.append(f"  条件付き書式を含むファイル: {len(cf_files)}件")
        lines.append(f"  印刷設定を含むファイル: {len(print_setting_files)}件")
        lines.append(f"  アウトラインを含むファイル: {len(outline_files)}件")
        lines.append(f"  オートフィルタを含むファイル: {len(filter_files)}件")

        layout_counter: Counter[str] = Counter(
            p.estimated_layout_type or "unknown" for p in xlsx_files
        )
        lines.append("  レイアウト推定の分布:")
        for layout, count in layout_counter.most_common():
            lines.append(f"    - {layout}: {count}件")
        lines.append("")

    # --- COM 変換が必要なファイル ---
    com_needed = [p for p in profiles if p.extension in (".doc", ".rtf", ".xls", ".ppt")]
    if com_needed:
        lines.append("■ COM 変換が必要なファイル（中身はまだ読めていない）")
        by_ext: dict[str, int] = {}
        for p in com_needed:
            by_ext[p.extension] = by_ext.get(p.extension, 0) + 1
        for ext, count in sorted(by_ext.items()):
            lines.append(f"  {ext}: {count}件")
        lines.append("  → これらは COM で変換後に再度プロファイリングすると詳細が取れます")
        lines.append("")

    # --- BAGLES ---
    bagles = [p for p in profiles if p.category == "BAGLES"]
    if bagles:
        lines.append("■ BAGLES ファイル")
        lines.append(f"  合計: {len(bagles)}件")
        for p in bagles:
            lines.append(f"  - {p.path} ({format_size(p.size_bytes)})")
        lines.append("")

    # --- エラー ---
    if summary["errors"]:
        lines.append(f"■ 読み込みエラー: {summary['error_count']}件")
        for e in summary["errors"]:
            lines.append(f"  - {e['path']}: {e['error']}")
        lines.append("")

    # --- フォルダ構造 ---
    folders: dict[str, list] = {}
    for p in profiles:
        folder = str(Path(p.path).parent)
        folders.setdefault(folder, []).append(p)
    lines.append(f"■ フォルダ数: {len(folders)}")
    lines.append("")
    for folder, items in sorted(folders.items()):
        exts = set(i.extension for i in items)
        lines.append(f"  {folder}/: {len(items)}ファイル (拡張子: {', '.join(sorted(exts))})")
    lines.append("")

    lines.append("=" * 60)
    lines.append("以上")
    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="文書プロファイリング: 指定フォルダ配下の全ファイルの構造情報を抽出する"
    )
    parser.add_argument("input_dir", help="スキャン対象のルートフォルダ")
    parser.add_argument(
        "--output",
        "-o",
        default="profile_report",
        help="出力ファイル名（拡張子なし）。.json と .txt を両方出力する (デフォルト: profile_report)",
    )
    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    if not input_dir.is_dir():
        print(f"エラー: {input_dir} はディレクトリではありません", file=sys.stderr)
        sys.exit(1)

    print(f"スキャン開始: {input_dir}")
    profiles = scan_directory(input_dir)
    print(f"スキャン完了: {len(profiles)} ファイル")

    summary = build_summary(profiles)

    # JSON レポート（詳細データ）
    json_path = Path(f"{args.output}.json")
    report = {
        "summary": summary,
        "files": [asdict(p) for p in profiles],
    }
    json_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"JSON レポート出力: {json_path}")

    # テキストレポート（口頭説明用）
    text_report = build_text_report(profiles, summary)
    txt_path = Path(f"{args.output}.txt")
    txt_path.write_text(text_report, encoding="utf-8")
    print(f"テキストレポート出力: {txt_path}")

    # コンソールにもテキストレポートを表示
    print("")
    print(text_report)


if __name__ == "__main__":
    main()
