"""Step2 構造抽出テスト

確認観点:
  - 中間 JSON が生成される
  - 疑似見出し検出（スタイル / フォントサイズ差 / ヒューリスティクス）
  - 表の抽出（セルテキスト・結合セル情報）
  - 変更履歴テーブル判定（全角スペース正規化含む）
  - doc_role 推定（spec_body / change_history / mixed / unknown）
  - 要素の出現順序が元文書と一致
"""

from pathlib import Path

import pytest

from src.config import PipelineConfig
from src.extractors.word import _group_shapes_as_flow, extract_docx
from src.models.intermediate import ShapeElement
from src.models.metadata import ProcessStatus
from tests.conftest import _create_dummy_image


class TestHeadingDetection:
    """疑似見出し検出のテスト"""

    def test_word_heading_style(self, simple_docx: Path, config: PipelineConfig):
        """Word 見出しスタイルで検出されること"""
        record, result = extract_docx(simple_docx, "simple.docx", ".docx", config)

        assert result.status in (ProcessStatus.SUCCESS, ProcessStatus.WARNING)
        elements = record.document["elements"]

        headings = [e for e in elements if e["type"] == "heading"]
        assert len(headings) >= 1, "見出しが1つ以上検出されること"

        # "第1章 概要" が level 1 で検出
        h1 = [h for h in headings if "概要" in h["content"]["text"]]
        assert len(h1) == 1
        assert h1[0]["content"]["level"] == 1
        assert h1[0]["content"]["detection_method"] == "style"

    def test_font_size_heading(self, font_size_heading_docx: Path, config: PipelineConfig):
        """フォントサイズ差で疑似見出しが検出されること (Oasys/Win 対応)"""
        record, result = extract_docx(
            font_size_heading_docx, "font_heading.docx", ".docx", config,
        )

        elements = record.document["elements"]
        headings = [e for e in elements if e["type"] == "heading"]
        assert len(headings) >= 2, "フォントサイズ差で2つ以上の見出しが検出されること"

        # 16pt → level 1
        h16 = [h for h in headings if "機能概要" in h["content"]["text"]]
        assert len(h16) == 1
        assert h16[0]["content"]["level"] == 1
        assert "font_size" in h16[0]["content"]["detection_method"]

        # 14pt → level 2
        h14 = [h for h in headings if "入力条件" in h["content"]["text"]]
        assert len(h14) == 1
        assert h14[0]["content"]["level"] == 2

        # 12pt → level 3
        h12 = [h for h in headings if "必須項目一覧" in h["content"]["text"]]
        assert len(h12) == 1
        assert h12[0]["content"]["level"] == 3

    def test_heuristic_heading(self, heuristic_heading_docx: Path, config: PipelineConfig):
        """短文 + 句点なしパターンでの疑似見出し検出"""
        record, result = extract_docx(
            heuristic_heading_docx, "heuristic.docx", ".docx", config,
        )

        elements = record.document["elements"]
        headings = [e for e in elements if e["type"] == "heading"]

        # "エラーコード一覧" と "入力チェック" が見出し候補
        heading_texts = [h["content"]["text"] for h in headings]
        assert "エラーコード一覧" in heading_texts
        assert "入力チェック" in heading_texts

        # 句点ありの本文は見出しにならない
        for h in headings:
            assert not h["content"]["text"].endswith("。")


class TestTableExtraction:
    """表の抽出テスト"""

    def test_basic_table(self, simple_docx: Path, config: PipelineConfig):
        """基本的な表のセルテキストが正しく抽出されること"""
        record, result = extract_docx(simple_docx, "simple.docx", ".docx", config)

        elements = record.document["elements"]
        tables = [e for e in elements if e["type"] == "table"]
        assert len(tables) == 1

        table = tables[0]["content"]
        rows = table["rows"]
        assert len(rows) == 3  # ヘッダー + 2データ行

        # ヘッダー行
        header_texts = [cell["text"] for cell in rows[0]]
        assert "項目" in header_texts
        assert "内容" in header_texts
        assert "備考" in header_texts

        # データ行
        assert rows[1][0]["text"] == "機能A"
        assert rows[1][1]["text"] == "入力チェック処理"

    def test_merged_cells(self, merged_cells_docx: Path, config: PipelineConfig):
        """結合セルが検出されること"""
        record, result = extract_docx(
            merged_cells_docx, "merged.docx", ".docx", config,
        )

        elements = record.document["elements"]
        tables = [e for e in elements if e["type"] == "table"]
        assert len(tables) >= 1

        table = tables[0]["content"]
        assert table["has_merged_cells"] is True
        assert table["confidence"] == "medium"


class TestChangeHistoryDetection:
    """変更履歴テーブル検出テスト"""

    def test_detect_with_fullwidth_spaces(
        self, change_history_docx: Path, config: PipelineConfig,
    ):
        """全角スペース入りのヘッダーでも変更履歴が検出されること"""
        record, result = extract_docx(
            change_history_docx, "change_history.docx", ".docx", config,
        )

        elements = record.document["elements"]
        ch_tables = [
            e for e in elements
            if e["type"] == "table"
            and e["content"].get("fallback_reason") == "change_history_table"
        ]
        assert len(ch_tables) == 1, "全角スペース正規化後に変更履歴テーブルが検出されること"

    def test_normal_table_not_detected(self, simple_docx: Path, config: PipelineConfig):
        """通常の表は変更履歴として検出されないこと"""
        record, result = extract_docx(simple_docx, "simple.docx", ".docx", config)

        elements = record.document["elements"]
        ch_tables = [
            e for e in elements
            if e["type"] == "table"
            and e["content"].get("fallback_reason") == "change_history_table"
        ]
        assert len(ch_tables) == 0


class TestDocRoleGuess:
    """文書の役割推定テスト"""

    def test_spec_body(self, simple_docx: Path, config: PipelineConfig):
        """仕様書本体の判定"""
        record, _ = extract_docx(simple_docx, "simple.docx", ".docx", config)
        assert record.metadata.doc_role_guess == "spec_body"

    def test_change_history(self, change_history_docx: Path, config: PipelineConfig):
        """変更履歴のみの判定"""
        record, _ = extract_docx(
            change_history_docx, "change_history.docx", ".docx", config,
        )
        assert record.metadata.doc_role_guess == "change_history"

    def test_mixed(self, mixed_docx: Path, config: PipelineConfig):
        """仕様書 + 変更履歴の混在判定"""
        record, _ = extract_docx(mixed_docx, "mixed.docx", ".docx", config)
        assert record.metadata.doc_role_guess == "mixed"

    def test_unknown(self, empty_docx: Path, config: PipelineConfig):
        """表なし → unknown"""
        record, _ = extract_docx(empty_docx, "empty.docx", ".docx", config)
        assert record.metadata.doc_role_guess == "unknown"


class TestMergedCellExtraction:
    """結合セル抽出の改善テスト（P1）"""

    def test_horizontal_merge_dedup(self, tmp_path: Path, config: PipelineConfig):
        """横結合セルが重複せず1回だけ抽出されること"""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        table = doc.add_table(rows=3, cols=4)
        for i, h in enumerate(["項目", "設定値", "デフォルト", "備考"]):
            table.rows[0].cells[i].text = h
        # 横結合行
        table.rows[1].cells[0].text = "■ 接続設定"
        table.rows[1].cells[0].merge(table.rows[1].cells[3])
        # 通常行
        table.rows[2].cells[0].text = "ホスト名"
        table.rows[2].cells[1].text = "db-server01"

        path = tmp_path / "hmerge.docx"
        doc.save(str(path))

        record, result = extract_docx(path, "hmerge.docx", ".docx", config)
        tables = [e for e in record.document["elements"] if e["type"] == "table"]
        assert len(tables) == 1

        rows = tables[0]["content"]["rows"]
        # 横結合行: 1セルのみ（colspan=4）
        banner_row = rows[1]
        assert len(banner_row) == 1, f"横結合行は1セルのみ: got {len(banner_row)}"
        assert banner_row[0]["text"] == "■ 接続設定"
        assert banner_row[0]["colspan"] == 4

    def test_vertical_merge_no_newline_dup(self, tmp_path: Path, config: PipelineConfig):
        """縦結合セルのテキストに改行重複が含まれないこと"""
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement

        doc = DocxDocument()
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "分類"
        table.rows[0].cells[1].text = "項目"
        table.rows[1].cells[0].text = "入力系"
        table.rows[1].cells[1].text = "ファイル"
        # vMerge restart
        tc_pr_1 = table.rows[1].cells[0]._tc.get_or_add_tcPr()
        vm_restart = OxmlElement("w:vMerge")
        vm_restart.set(_qn("w:val"), "restart")
        tc_pr_1.append(vm_restart)
        # vMerge continue
        table.rows[2].cells[1].text = "DB"
        tc_pr_2 = table.rows[2].cells[0]._tc.get_or_add_tcPr()
        vm_continue = OxmlElement("w:vMerge")
        tc_pr_2.append(vm_continue)

        path = tmp_path / "vmerge.docx"
        doc.save(str(path))

        record, result = extract_docx(path, "vmerge.docx", ".docx", config)
        tables = [e for e in record.document["elements"] if e["type"] == "table"]

        # セルテキストに \n が含まれないこと
        for row in tables[0]["content"]["rows"]:
            for cell in row:
                assert "\n" not in cell["text"], f"セルテキストに改行: {cell['text']}"


class TestElementOrder:
    """要素の出現順序テスト"""

    def test_paragraph_table_order(self, simple_docx: Path, config: PipelineConfig):
        """段落→表→段落の順序が元文書と一致すること"""
        record, _ = extract_docx(simple_docx, "simple.docx", ".docx", config)

        elements = record.document["elements"]
        types = [e["type"] for e in elements]

        # heading → paragraph → heading → paragraph → table → heading → paragraph
        # の順序であること（空段落除去により多少変わりうる）
        assert types[0] == "heading"  # 第1章 概要

        # 表がどこかに存在
        assert "table" in types

        # source_index が単調増加
        indices = [e["source_index"] for e in elements]
        assert indices == sorted(indices), "source_index が単調増加であること"


class TestImageCaptionDetection:
    """画像キャプション検出テスト（P2）"""

    def test_figure_caption_merged_to_image(self, tmp_path: Path, config: PipelineConfig):
        """画像直後の「図: XXX」が ImageElement.description に統合されること"""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("以下にシステム構成図を示す。")
        image_path = tmp_path / "dummy.png"
        _create_dummy_image(image_path)
        doc.add_picture(str(image_path))
        doc.add_paragraph("図: システム構成図")
        doc.add_paragraph("本文が続く。")

        path = tmp_path / "img_caption.docx"
        doc.save(str(path))

        record, result = extract_docx(path, "img_caption.docx", ".docx", config)

        assert result.status in (ProcessStatus.SUCCESS, ProcessStatus.WARNING)
        elements = record.document["elements"]

        images = [e for e in elements if e["type"] == "image"]
        assert len(images) == 1
        assert images[0]["content"]["description"] == "システム構成図"

        headings = [e for e in elements if e["type"] == "heading"]
        caption_headings = [
            h for h in headings if "システム構成図" in h["content"]["text"]
        ]
        assert len(caption_headings) == 0

    def test_non_caption_not_merged(self, tmp_path: Path, config: PipelineConfig):
        """画像直後でも図キャプションでない段落は統合されないこと"""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        image_path = tmp_path / "dummy.png"
        _create_dummy_image(image_path)
        doc.add_picture(str(image_path))
        doc.add_paragraph("この画像は参考資料である。")

        path = tmp_path / "img_no_caption.docx"
        doc.save(str(path))

        record, _ = extract_docx(path, "img_no_caption.docx", ".docx", config)
        elements = record.document["elements"]

        images = [e for e in elements if e["type"] == "image"]
        assert len(images) == 1
        assert images[0]["content"]["description"] == ""

        paragraphs = [e for e in elements if e["type"] == "paragraph"]
        assert any(
            p["content"]["text"] == "この画像は参考資料である。"
            for p in paragraphs
        )

    def test_figure_caption_with_number(self, tmp_path: Path, config: PipelineConfig):
        """「図1 画面遷移図」形式のキャプションが統合されること"""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        image_path = tmp_path / "dummy.png"
        _create_dummy_image(image_path)
        doc.add_picture(str(image_path))
        doc.add_paragraph("図1 画面遷移図")

        path = tmp_path / "img_numbered.docx"
        doc.save(str(path))

        record, _ = extract_docx(path, "img_numbered.docx", ".docx", config)
        elements = record.document["elements"]

        images = [e for e in elements if e["type"] == "image"]
        assert len(images) == 1
        assert "画面遷移図" in images[0]["content"]["description"]


class TestFlowGrouping:
    """図形フローグルーピングテスト（P3）"""

    def test_arrow_text_not_heading(self, tmp_path: Path, config: PipelineConfig):
        """矢印テキストが見出しに昇格しないこと"""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("→ → → → →")
        doc.add_paragraph("正常系: → → → →")
        doc.add_paragraph("↓ の矢印で接続")
        doc.add_paragraph("エラー処理について")

        path = tmp_path / "arrows.docx"
        doc.save(str(path))

        record, _ = extract_docx(path, "arrows.docx", ".docx", config)
        elements = record.document["elements"]

        headings = [e for e in elements if e["type"] == "heading"]
        heading_texts = [h["content"]["text"] for h in headings]

        assert "→ → → → →" not in heading_texts
        assert "正常系: → → → →" not in heading_texts
        assert "↓ の矢印で接続" not in heading_texts
        assert "エラー処理について" in heading_texts

    def test_group_shapes_as_workflow_sorted_by_position(self):
        """3個以上のテキスト図形が位置順に workflow へまとめられること"""
        shapes = [
            ShapeElement(
                shape_type="vml_textbox",
                texts=["部長", "承認判断"],
                left_pt=200,
                top_pt=0,
            ),
            ShapeElement(
                shape_type="vml_textbox",
                texts=["申請者", "申請書作成"],
                left_pt=0,
                top_pt=0,
            ),
            ShapeElement(
                shape_type="vml_textbox",
                texts=["上長", "内容確認"],
                left_pt=100,
                top_pt=0,
            ),
        ]

        grouped = _group_shapes_as_flow(shapes)

        assert len(grouped) == 1
        assert grouped[0].shape_type == "workflow"
        assert grouped[0].texts == [
            "申請者 / 申請書作成",
            "上長 / 内容確認",
            "部長 / 承認判断",
        ]


class TestHeadingPrecision:
    """見出し検出精度テスト（P4）"""

    def test_section_number_overrides_font_size_level(
        self, tmp_path: Path, config: PipelineConfig,
    ):
        """section 番号で font-size 見出しのレベルが補正されること"""
        from docx import Document as DocxDocument
        from docx.shared import Pt

        doc = DocxDocument()
        doc.add_heading("4. 機能詳細", level=2)

        p1 = doc.add_paragraph()
        run1 = p1.add_run("4.1 入力チェック機能")
        run1.font.size = Pt(14)
        doc.add_paragraph("入力データの妥当性を検証する。")

        p2 = doc.add_paragraph()
        run2 = p2.add_run("4.2 データ出力機能")
        run2.font.size = Pt(14)

        path = tmp_path / "section_num.docx"
        doc.save(str(path))

        record, _ = extract_docx(path, "section_num.docx", ".docx", config)
        elements = record.document["elements"]

        headings = [e for e in elements if e["type"] == "heading"]
        h_parent = [h for h in headings if "機能詳細" in h["content"]["text"]]
        assert h_parent[0]["content"]["level"] == 2

        h_child = [h for h in headings if "入力チェック" in h["content"]["text"]]
        assert h_child[0]["content"]["level"] == 3

    def test_table_caption_not_heading(self, tmp_path: Path, config: PipelineConfig):
        """「表N: XXX」パターンがヒューリスティクス見出しにならないこと"""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("表1: ユーザー管理テーブル")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "ID"
        table.rows[0].cells[1].text = "名前"
        table.rows[1].cells[0].text = "1"
        table.rows[1].cells[1].text = "田中"

        path = tmp_path / "table_caption.docx"
        doc.save(str(path))

        record, _ = extract_docx(path, "table_caption.docx", ".docx", config)
        elements = record.document["elements"]

        headings = [e for e in elements if e["type"] == "heading"]
        heading_texts = [h["content"]["text"] for h in headings]
        assert "表1: ユーザー管理テーブル" not in heading_texts

    def test_figure_caption_standalone_not_heading(
        self, tmp_path: Path, config: PipelineConfig,
    ):
        """画像直後でない「図: XXX」もヒューリスティクス見出しにならないこと"""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("前のセクション。")
        doc.add_paragraph("")
        doc.add_paragraph("図: 概念図")

        path = tmp_path / "standalone_fig_caption.docx"
        doc.save(str(path))

        record, _ = extract_docx(
            path, "standalone_fig_caption.docx", ".docx", config,
        )
        elements = record.document["elements"]

        headings = [e for e in elements if e["type"] == "heading"]
        heading_texts = [h["content"]["text"] for h in headings]
        assert "図: 概念図" not in heading_texts

    def test_chapter_heading_level(self, tmp_path: Path, config: PipelineConfig):
        """「第N章」パターンが L2 で検出されること"""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("第1章 システム概要")
        doc.add_paragraph("本章ではシステムの概要を述べる。")

        path = tmp_path / "chapter.docx"
        doc.save(str(path))

        record, _ = extract_docx(path, "chapter.docx", ".docx", config)
        elements = record.document["elements"]

        headings = [e for e in elements if e["type"] == "heading"]
        h = [heading for heading in headings if "システム概要" in heading["content"]["text"]]
        assert len(h) == 1
        assert h[0]["content"]["level"] == 2


class TestErrorHandling:
    """エラーハンドリングテスト"""

    def test_nonexistent_file(self, tmp_path: Path, config: PipelineConfig):
        """存在しないファイルでエラーが返ること"""
        fake_path = tmp_path / "nonexistent.docx"
        record, result = extract_docx(fake_path, "nonexistent.docx", ".docx", config)

        assert result.status == ProcessStatus.ERROR
        assert record.document == {}
