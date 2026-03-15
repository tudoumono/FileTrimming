"""テスト用フィクスチャ

python-docx / openpyxl で各種パターンの .docx / .xlsx を動的に生成する。
実データは使えないため、§14 の調査結果に基づいた再現ファイルを作る。
"""

from __future__ import annotations

import struct
import zlib
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.comments import Comment

from src.config import PipelineConfig


def _create_dummy_image(path: Path) -> None:
    """テスト用 1x1 PNG を作成する。"""
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    ihdr_crc = zlib.crc32(b"IHDR" + ihdr_data)
    ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + struct.pack(">I", ihdr_crc)

    raw = b"\x00\xff\xff\xff"
    idat_data = zlib.compress(raw)
    idat_crc = zlib.crc32(b"IDAT" + idat_data)
    idat = (
        struct.pack(">I", len(idat_data))
        + b"IDAT"
        + idat_data
        + struct.pack(">I", idat_crc)
    )

    iend_crc = zlib.crc32(b"IEND")
    iend = struct.pack(">I", 0) + b"IEND" + struct.pack(">I", iend_crc)
    path.write_bytes(sig + ihdr + idat + iend)


@pytest.fixture
def config(tmp_path: Path) -> PipelineConfig:
    """テスト用パイプライン設定（tmp_path ベース）"""
    return PipelineConfig(
        input_dir=tmp_path / "input",
        intermediate_base=tmp_path / "intermediate",
        output_base=tmp_path / "output",
        run_id="test_run",
    )


@pytest.fixture
def simple_docx(tmp_path: Path) -> Path:
    """見出し + 段落 + 表を含む基本的な .docx"""
    doc = Document()

    # 見出し（Word 見出しスタイル）
    doc.add_heading("第1章 概要", level=1)
    doc.add_paragraph("この文書はテスト用の仕様書です。")

    doc.add_heading("1.1 目的", level=2)
    doc.add_paragraph("テストパイプラインの動作確認を目的とする。")

    # 表
    table = doc.add_table(rows=3, cols=3)
    headers = ["項目", "内容", "備考"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "機能A"
    table.rows[1].cells[1].text = "入力チェック処理"
    table.rows[1].cells[2].text = "必須"
    table.rows[2].cells[0].text = "機能B"
    table.rows[2].cells[1].text = "データ更新処理"
    table.rows[2].cells[2].text = ""

    doc.add_heading("1.2 制約事項", level=2)
    doc.add_paragraph("オフライン環境で動作すること。")

    path = tmp_path / "simple.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def change_history_docx(tmp_path: Path) -> Path:
    """変更履歴テーブルを含む .docx

    §14 の検出条件: 1行目に「ページ」「種別」「年月」「記事」のうち3個以上
    全角スペース入りのケースも再現する。
    """
    doc = Document()

    doc.add_paragraph("変更履歴")

    # 変更履歴テーブル（全角スペース入り）
    table = doc.add_table(rows=4, cols=4)
    # ヘッダー行: 全角スペースを含む（§14 で確認された実パターン）
    table.rows[0].cells[0].text = "ページ"
    table.rows[0].cells[1].text = "種別"
    table.rows[0].cells[2].text = "年\u3000月"      # 全角スペース入り
    table.rows[0].cells[3].text = "記\u3000\u3000事"  # 全角スペース2個

    table.rows[1].cells[0].text = "3"
    table.rows[1].cells[1].text = "追加"
    table.rows[1].cells[2].text = "2025/01"
    table.rows[1].cells[3].text = "初版作成"

    table.rows[2].cells[0].text = "5"
    table.rows[2].cells[1].text = "修正"
    table.rows[2].cells[2].text = "2025/03"
    table.rows[2].cells[3].text = "エラー処理追加"

    table.rows[3].cells[0].text = "10"
    table.rows[3].cells[1].text = "削除"
    table.rows[3].cells[2].text = "2025/06"
    table.rows[3].cells[3].text = "旧機能削除"

    path = tmp_path / "change_history.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def font_size_heading_docx(tmp_path: Path) -> Path:
    """フォントサイズ差による疑似見出しを含む .docx

    §14: Oasys/Win スタイルでは見出しスタイルが使われないため、
    フォントサイズ差で見出し検出する必要がある。
    """
    doc = Document()

    # 大きいフォント = 見出し候補
    p1 = doc.add_paragraph()
    run1 = p1.add_run("機能概要")
    run1.font.size = Pt(16)

    # 本文サイズ
    p2 = doc.add_paragraph()
    run2 = p2.add_run("この機能は入力データの整合性を検証する。")
    run2.font.size = Pt(10.5)

    # 中サイズ = サブ見出し候補
    p3 = doc.add_paragraph()
    run3 = p3.add_run("入力条件")
    run3.font.size = Pt(14)

    # 本文サイズ
    p4 = doc.add_paragraph()
    run4 = p4.add_run("入力データは CSV 形式であること。")
    run4.font.size = Pt(10.5)

    # さらに小さい見出し候補
    p5 = doc.add_paragraph()
    run5 = p5.add_run("必須項目一覧")
    run5.font.size = Pt(12)

    p6 = doc.add_paragraph()
    run6 = p6.add_run("顧客番号、氏名、住所は必須とする。")
    run6.font.size = Pt(10.5)

    path = tmp_path / "font_heading.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def merged_cells_docx(tmp_path: Path) -> Path:
    """結合セルを含む表の .docx"""
    doc = Document()
    doc.add_paragraph("結合セルテスト")

    table = doc.add_table(rows=4, cols=3)
    table.rows[0].cells[0].text = "分類"
    table.rows[0].cells[1].text = "項目"
    table.rows[0].cells[2].text = "値"

    # 縦結合: 1行目と2行目の「分類」列を結合
    table.rows[1].cells[0].text = "入力系"
    table.rows[1].cells[1].text = "ファイル名"
    table.rows[1].cells[2].text = "data.csv"

    table.rows[2].cells[0].text = "入力系"  # 結合対象（同じテキスト）
    table.rows[2].cells[1].text = "文字コード"
    table.rows[2].cells[2].text = "UTF-8"

    # 横結合
    table.rows[3].cells[0].text = "備考"
    # 横結合は python-docx の merge で実現
    table.rows[3].cells[0].merge(table.rows[3].cells[2])

    path = tmp_path / "merged.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def shapes_docx(tmp_path: Path) -> Path:
    """図形（テキストボックス）を含む .docx

    python-docx では浮動図形の直接追加が困難なため、
    InlineShape を使って最低限のテストを行う。
    """
    doc = Document()
    doc.add_paragraph("フロー図の説明")
    # python-docx で浮動図形は追加しにくいが、段落は追加できる
    doc.add_paragraph("開始 → 入力チェック → 処理実行 → 終了")
    doc.add_paragraph("分岐: エラー時 → エラー処理 → 終了")

    path = tmp_path / "shapes.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def mixed_docx(tmp_path: Path) -> Path:
    """仕様書本体 + 変更履歴が混在する .docx"""
    doc = Document()

    # 仕様書部分
    doc.add_heading("機能仕様書", level=1)
    doc.add_paragraph("本文書は機能Aの仕様を定義する。")

    # 通常の表
    spec_table = doc.add_table(rows=2, cols=2)
    spec_table.rows[0].cells[0].text = "項目"
    spec_table.rows[0].cells[1].text = "内容"
    spec_table.rows[1].cells[0].text = "入力"
    spec_table.rows[1].cells[1].text = "CSV ファイル"

    doc.add_paragraph("")

    # 変更履歴テーブル
    ch_table = doc.add_table(rows=2, cols=4)
    ch_table.rows[0].cells[0].text = "ページ"
    ch_table.rows[0].cells[1].text = "種別"
    ch_table.rows[0].cells[2].text = "年月"
    ch_table.rows[0].cells[3].text = "記事"
    ch_table.rows[1].cells[0].text = "1"
    ch_table.rows[1].cells[1].text = "新規"
    ch_table.rows[1].cells[2].text = "2025/01"
    ch_table.rows[1].cells[3].text = "初版"

    path = tmp_path / "mixed.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def empty_docx(tmp_path: Path) -> Path:
    """空の .docx"""
    doc = Document()
    path = tmp_path / "empty.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def heuristic_heading_docx(tmp_path: Path) -> Path:
    """短文ヒューリスティクスによる疑似見出し検出テスト用 .docx

    フォントサイズ指定なし・見出しスタイルなし・30文字以下・句点なし
    → heuristic:short_no_period で検出されるべき。
    """
    doc = Document()
    doc.add_paragraph("エラーコード一覧")  # 見出し候補（短文・句点なし）
    doc.add_paragraph("以下にエラーコードの一覧を示す。")  # 本文（句点あり）
    doc.add_paragraph("入力チェック")  # 見出し候補
    doc.add_paragraph("入力されたデータが不正な場合はエラーとする。")  # 本文

    path = tmp_path / "heuristic.docx"
    doc.save(str(path))
    return path


# ===========================================================================
# Excel テスト用フィクスチャ
# ===========================================================================

@pytest.fixture
def simple_xlsx(tmp_path: Path) -> Path:
    """基本的なデータテーブルを含む .xlsx"""
    wb = Workbook()
    ws = wb.active
    ws.title = "エラーコード一覧"

    # ヘッダー
    ws.append(["コード", "メッセージ", "対処方法"])
    ws.append(["E001", "入力値が不正です", "入力形式を確認してください"])
    ws.append(["E002", "接続タイムアウト", "ネットワーク状態を確認してください"])
    ws.append(["E003", "ファイルが見つかりません", "パスを確認してください"])

    path = tmp_path / "simple.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture
def multi_sheet_xlsx(tmp_path: Path) -> Path:
    """複数シートを含む .xlsx"""
    wb = Workbook()

    # シート1: 機能一覧
    ws1 = wb.active
    ws1.title = "機能一覧"
    ws1.append(["機能ID", "機能名", "ステータス"])
    ws1.append(["F001", "ログイン", "完了"])
    ws1.append(["F002", "データ検索", "開発中"])

    # シート2: エラーコード
    ws2 = wb.create_sheet("エラーコード")
    ws2.append(["コード", "内容"])
    ws2.append(["E001", "認証エラー"])

    path = tmp_path / "multi_sheet.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture
def merged_cells_xlsx(tmp_path: Path) -> Path:
    """結合セルを含む .xlsx"""
    wb = Workbook()
    ws = wb.active
    ws.title = "設定一覧"

    # ヘッダー行（横結合）
    ws["A1"] = "分類"
    ws["B1"] = "項目"
    ws["C1"] = "値"

    # 縦結合: A2:A3 を結合
    ws["A2"] = "接続設定"
    ws.merge_cells("A2:A3")
    ws["B2"] = "ホスト名"
    ws["C2"] = "db-server01"
    ws["B3"] = "ポート"
    ws["C3"] = "5432"

    # 横結合: A4:C4 を結合（セクション区切り）
    ws["A4"] = "■ 認証設定"
    ws.merge_cells("A4:C4")

    ws["A5"] = "認証設定"
    ws["B5"] = "ユーザー名"
    ws["C5"] = "admin"

    path = tmp_path / "merged.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture
def hidden_sheet_xlsx(tmp_path: Path) -> Path:
    """非表示シートを含む .xlsx"""
    wb = Workbook()

    ws1 = wb.active
    ws1.title = "表示シート"
    ws1.append(["データ", "値"])
    ws1.append(["A", "1"])

    ws2 = wb.create_sheet("マスタ")
    ws2.append(["マスタデータ"])
    ws2.append(["非表示の内容"])
    ws2.sheet_state = "hidden"

    path = tmp_path / "hidden.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture
def comments_xlsx(tmp_path: Path) -> Path:
    """コメント付きセルを含む .xlsx"""
    wb = Workbook()
    ws = wb.active
    ws.title = "レビュー"

    ws.append(["項目", "値", "備考"])
    ws["A2"] = "顧客ID"
    ws["B2"] = "C001"
    ws["B2"].comment = Comment("外部IFとの整合確認が必要。", "レビュアー")
    ws["A3"] = "契約日"
    ws["B3"] = "2025/01/15"

    path = tmp_path / "comments.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture
def empty_xlsx(tmp_path: Path) -> Path:
    """空の .xlsx"""
    wb = Workbook()
    ws = wb.active
    ws.title = "空シート"
    # 何もデータを入れない

    path = tmp_path / "empty.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture
def large_xlsx(tmp_path: Path) -> Path:
    """大きなデータシートを含む .xlsx（警告閾値テスト用）"""
    wb = Workbook()
    ws = wb.active
    ws.title = "取引明細"

    headers = ["取引ID", "日付", "顧客名", "金額"]
    ws.append(headers)
    for i in range(600):  # 500行超 → 警告
        ws.append([f"T{i+1:04d}", "2025/01/01", f"顧客{i+1}", (i + 1) * 1000])

    path = tmp_path / "large.xlsx"
    wb.save(str(path))
    return path
